"""
Base de Conhecimento (RAG) — aprendizado a partir de documentos de referência.

O usuário envia arquivos (PDF, DOCX, TXT/MD) de leis, acórdãos,
entendimentos dos Tribunais de Contas, processos anteriores e modelos.
Cada arquivo é dividido em trechos (chunks), que recebem embeddings do
Gemini e são armazenados no Supabase (pgvector). Na geração de cada
documento, os trechos mais relevantes são recuperados e injetados no
prompt como fundamentação.

Estratégia de busca:
  1. Vetorial (pgvector + embeddings Gemini) — quando há chave de API;
  2. Textual em português (tsvector/websearch) — fallback automático.
"""

import io
import re

import streamlit as st

from . import db
from .config import (
    EMBEDDING_DIMENSOES,
    EMBEDDING_MODEL,
    RAG_CHUNK_SOBREPOSICAO,
    RAG_CHUNK_TAMANHO,
    RAG_TOP_K,
)

CATEGORIAS = {
    "lei": "Lei / Norma",
    "acordao": "Acórdão (TCU/TCE)",
    "entendimento": "Entendimento / Orientação de TC",
    "processo_anterior": "Processo anterior realizado",
    "modelo": "Modelo / Minuta padrão (AGU etc.)",
    "outro": "Outro",
}

# Hierarquia da fonte recuperada (P1): o que cada categoria PODE
# sustentar. Legislação e jurisprudência NÃO se confundem — acórdão e
# entendimento de Tribunal de Contas orientam a interpretação e o
# controle, mas não são a norma. Processo anterior é molde de estrutura
# e linguagem, nunca prova do direito vigente.
LEGISLACAO = ("lei",)
CONTROLE = ("acordao", "entendimento")
MOLDES = ("processo_anterior", "modelo")
NORMATIVAS = LEGISLACAO + CONTROLE   # compatibilidade: fontes jurídicas

_PAPEL_DA_FONTE = {
    "lei": "legislação/regulamento — fundamenta diretamente a cláusula",
    "acordao": "jurisprudência de controle — orienta a interpretação; "
               "não substitui a norma",
    "entendimento": "orientação de órgão de controle — orienta a "
                    "interpretação; não substitui a norma",
    "processo_anterior": "processo anterior — apenas estrutura e "
                         "linguagem; NÃO fundamenta",
    "modelo": "modelo/minuta padrão — apenas estrutura e linguagem; "
              "NÃO fundamenta",
}

# ---------------------------------------------------------------------------
# Recuperação TEMÁTICA (P1)
#
# A consulta única (documento + objeto + justificativa) recupera contexto
# geral, mas não garante que uma afirmação jurídica específica —
# "vigência da ata", "pagamento", "sanções" — tenha sido sustentada pelo
# dispositivo correspondente. Foi assim que artigos errados entraram nos
# documentos (pregão fundado no art. 109, pagamento no art. 98).
#
# A estratégia é uma LISTA CONTROLADA de temas, filtrada pelo documento e
# pelos fatos da contratação, com orçamento fixo de buscas. Todas as
# consultas viajam em UMA única chamada de embeddings (lote) e cada busca
# usa top-k reduzido — o custo fica próximo do fluxo atual.
# ---------------------------------------------------------------------------
# chave: (rótulo, termos de busca, condição)
# condição: None (sempre aplicável) ou gatilho estrutural de `_gatilhos`
TEMAS_JURIDICOS: dict[str, tuple[str, str, str | None]] = {
    "modalidade": (
        "Modalidade e critério de julgamento",
        "modalidade de licitação pregão concorrência critério de "
        "julgamento menor preço maior desconto", None),
    "srp": (
        "Sistema de Registro de Preços",
        "sistema de registro de preços ata de registro de preços vigência "
        "da ata adesão órgão gerenciador cadastro de reserva", "srp"),
    "parcelamento": (
        "Parcelamento e adjudicação por item ou lote",
        "parcelamento do objeto adjudicação por item lote divisibilidade "
        "economia de escala", None),
    "requisitos": (
        "Requisitos técnicos e habilitação",
        "requisitos da contratação especificação técnica habilitação "
        "jurídica fiscal técnica econômico-financeira qualificação", None),
    "execucao_recebimento": (
        "Execução, recebimento e aceitação do objeto",
        "prazo de entrega execução do objeto recebimento provisório "
        "definitivo aceitação do objeto", None),
    "pagamento": (
        "Pagamento e liquidação",
        "pagamento liquidação da despesa ordem cronológica prazo de "
        "pagamento nota fiscal atesto", None),
    "reajuste": (
        "Reajuste e repactuação de preços",
        "reajuste de preços índice repactuação equilíbrio "
        "econômico-financeiro revisão", None),
    "garantia": (
        "Garantia contratual",
        "garantia contratual caução seguro-garantia fiança bancária "
        "percentual da garantia", "garantia"),
    "gestao_fiscalizacao": (
        "Gestão e fiscalização do contrato",
        "gestor do contrato fiscal do contrato fiscalização registro de "
        "ocorrências", None),
    "sancoes": (
        "Infrações e sanções administrativas",
        "infrações administrativas sanções advertência multa impedimento "
        "de licitar declaração de inidoneidade", None),
    "recursos": (
        "Impugnações e recursos",
        "impugnação ao edital pedido de esclarecimento recurso "
        "administrativo prazo recursal", None),
    "me_epp": (
        "Tratamento favorecido a ME/EPP",
        "microempresa empresa de pequeno porte tratamento favorecido "
        "empate ficto regularidade fiscal", None),
    "protecao_dados": (
        "Proteção de dados e segurança da informação",
        "proteção de dados pessoais LGPD segurança da informação "
        "confidencialidade níveis de serviço", "dados"),
    "necessidade": (
        "Necessidade, estudo técnico preliminar e planejamento",
        "estudo técnico preliminar descrição da necessidade levantamento "
        "de soluções alternativas viabilidade planejamento da contratação",
        None),
}

# TEMAS NÚCLEO: as matérias que o documento SEMPRE decide. Recuperação
# garantida — nenhuma delas pode ficar sem consulta por disputa de vaga
# com um tema condicional (foi assim que 'sanções' e 'pagamento' ficaram
# de fora do TR na primeira versão).
TEMAS_NUCLEO: dict[str, tuple[str, ...]] = {
    "dfd": ("necessidade",),
    "etp": ("necessidade", "requisitos", "parcelamento", "modalidade"),
    "tr": ("execucao_recebimento", "pagamento", "sancoes",
           "gestao_fiscalizacao"),
    "edital": ("modalidade", "requisitos", "sancoes", "recursos"),
}

# TEMAS COMPLEMENTARES: entram conforme o objeto/modelagem. Os que têm
# gatilho satisfeito vêm primeiro; os demais ocupam a folga do orçamento.
TEMAS_COMPLEMENTARES: dict[str, tuple[str, ...]] = {
    "dfd": ("srp", "parcelamento"),
    "etp": ("srp", "protecao_dados", "reajuste", "me_epp"),
    "tr": ("srp", "garantia", "protecao_dados", "reajuste", "requisitos"),
    "edital": ("srp", "garantia", "protecao_dados", "me_epp",
               "parcelamento"),
}

# Orçamento de recuperação (custo/latência sob controle): no máximo
# 1 (geral) + 4 (núcleo) + 3 (complementares) = 8 buscas, sempre com UMA
# única chamada de embeddings em lote.
MAX_TEMAS_NUCLEO = 4
MAX_TEMAS_COMPLEMENTARES = 3
MAX_TEMAS = MAX_TEMAS_NUCLEO + MAX_TEMAS_COMPLEMENTARES
TOP_K_TEMA = 3              # trechos por busca temática
MAX_CHUNKS_PROMPT = 10      # teto de trechos enviados à IA

# Piso de relevância por modo de busca. Escalas diferentes: similaridade
# de cosseno (0..1) no vetorial; ts_rank (tipicamente < 0,1) no textual.
# Conservadores e sobrescrevíveis em config_app (`rag_piso_vetorial` /
# `rag_piso_textual`) — sem descartar a busca textual, que é o fallback
# quando não há embeddings.
PISO_VETORIAL_PADRAO = 0.20
PISO_TEXTUAL_PADRAO = 0.01

# O processo pede o tema de proteção de dados mesmo sem ser software
# (ex.: serviço que manipula base de dados de pacientes).
_TERMOS_DADOS_NO_PROCESSO = re.compile(
    r"dados\s+pessoais|lgpd|prote[çc][ãa]o\s+de\s+dados|sigilo|"
    r"seguran[çc]a\s+da\s+informa[çc][ãa]o|confidencialidade",
    re.IGNORECASE)


class ErroRAG(Exception):
    """Erro da base de conhecimento com mensagem amigável."""


# ---------------------------------------------------------------------------
# Extração de texto dos arquivos enviados
# ---------------------------------------------------------------------------
def extrair_texto(nome_arquivo: str, dados: bytes) -> str:
    """Extrai texto de PDF, DOCX, TXT ou MD. Levanta ErroRAG se não suportado."""
    extensao = nome_arquivo.rsplit(".", 1)[-1].lower() if "." in nome_arquivo else ""
    try:
        if extensao == "pdf":
            from pypdf import PdfReader

            leitor = PdfReader(io.BytesIO(dados))
            paginas = [pagina.extract_text() or "" for pagina in leitor.pages]
            texto = "\n".join(paginas)
        elif extensao == "docx":
            from docx import Document

            documento = Document(io.BytesIO(dados))
            partes = [p.text for p in documento.paragraphs]
            for tabela in documento.tables:
                for linha in tabela.rows:
                    partes.append(" | ".join(c.text for c in linha.cells))
            texto = "\n".join(partes)
        elif extensao in ("txt", "md"):
            texto = dados.decode("utf-8", errors="replace")
        else:
            raise ErroRAG(
                f"Formato .{extensao or '?'} não suportado — envie PDF, DOCX, TXT ou MD."
            )
    except ErroRAG:
        raise
    except Exception as exc:  # noqa: BLE001 — arquivo corrompido/ilegível
        raise ErroRAG(f"Não foi possível ler '{nome_arquivo}': {exc}") from exc

    texto = re.sub(r"[ \t]+", " ", texto)
    texto = re.sub(r"\n{3,}", "\n\n", texto).strip()
    if len(texto) < 50:
        raise ErroRAG(
            f"'{nome_arquivo}' não contém texto extraível (PDF digitalizado sem OCR?)."
        )
    return texto


def dividir_em_chunks(
    texto: str,
    tamanho: int = RAG_CHUNK_TAMANHO,
    sobreposicao: int = RAG_CHUNK_SOBREPOSICAO,
) -> list[str]:
    """
    Divide o texto em trechos de ~`tamanho` caracteres com sobreposição,
    preferindo quebrar em fim de parágrafo ou sentença para não cortar
    dispositivos legais no meio.
    """
    chunks: list[str] = []
    inicio = 0
    while inicio < len(texto):
        fim = min(inicio + tamanho, len(texto))
        if fim < len(texto):
            janela = texto[inicio:fim]
            # tenta quebrar no último parágrafo; senão na última sentença
            corte = max(janela.rfind("\n\n"), janela.rfind(". "))
            if corte > tamanho // 2:
                fim = inicio + corte + 1
        trecho = texto[inicio:fim].strip()
        if trecho:
            chunks.append(trecho)
        if fim >= len(texto):
            break
        inicio = max(fim - sobreposicao, inicio + 1)
    return chunks


# ---------------------------------------------------------------------------
# Embeddings (Gemini) — opcionais; sem eles a busca textual assume
# ---------------------------------------------------------------------------
def _gerar_embeddings(textos: list[str], para_consulta: bool) -> list[list[float]] | None:
    """
    Retorna embeddings (768 dims) ou None se não houver chave de API.

    Provedor segue o motor principal: OpenAI (text-embedding-3-small com
    dimensions=768) quando há chave; senão Gemini. IMPORTANTE: indexação e
    consulta precisam do MESMO provedor — se você trocar de provedor com a
    base já populada, reindexe os arquivos (os espaços vetoriais são
    incompatíveis entre si).
    """
    from .config import OPENAI_EMBEDDING_MODEL
    from .llm import obter_api_key, obter_openai_key

    chave_openai = obter_openai_key()
    chave_gemini = obter_api_key()
    try:
        if chave_openai:
            from openai import OpenAI

            cliente = OpenAI(api_key=chave_openai, timeout=60, max_retries=1)
            resposta = cliente.embeddings.create(
                model=OPENAI_EMBEDDING_MODEL,
                input=textos,
                dimensions=EMBEDDING_DIMENSOES,
            )
            return [item.embedding for item in resposta.data]
        if chave_gemini:
            from google import genai
            from google.genai import types

            cliente = genai.Client(api_key=chave_gemini)
            resposta = cliente.models.embed_content(
                model=EMBEDDING_MODEL,
                contents=textos,
                config=types.EmbedContentConfig(
                    task_type="RETRIEVAL_QUERY" if para_consulta else "RETRIEVAL_DOCUMENT",
                    output_dimensionality=EMBEDDING_DIMENSOES,
                ),
            )
            return [list(e.values) for e in resposta.embeddings]
        return None
    except Exception as exc:  # noqa: BLE001
        # Falha de embedding não deve impedir a indexação: busca textual assume
        st.warning(f"Embeddings indisponíveis ({exc}); usando busca textual.")
        return None


# ---------------------------------------------------------------------------
# Indexação e gestão da biblioteca
# ---------------------------------------------------------------------------
def indexar_arquivo(nome_arquivo: str, titulo: str, categoria: str, dados: bytes) -> int:
    """Extrai, divide, gera embeddings e grava o documento. Retorna nº de chunks."""
    if not db.disponivel():
        raise ErroRAG(
            "A Base de Conhecimento exige o Supabase configurado "
            "(SUPABASE_URL e SUPABASE_KEY em .streamlit/secrets.toml)."
        )
    texto = extrair_texto(nome_arquivo, dados)
    chunks = dividir_em_chunks(texto)
    if not chunks:
        raise ErroRAG(f"'{nome_arquivo}' não gerou trechos indexáveis.")

    embeddings = _gerar_embeddings(chunks, para_consulta=False)

    try:
        cliente = db._cliente()  # noqa: SLF001 — reuso interno do cliente único
        doc = (
            cliente.table("documentos_referencia")
            .insert(
                {
                    "titulo": titulo or nome_arquivo,
                    "categoria": categoria,
                    "nome_arquivo": nome_arquivo,
                    "total_chunks": len(chunks),
                }
            )
            .execute()
        ).data[0]

        registros = [
            {
                "documento_id": doc["id"],
                "ordem": i,
                "conteudo": trecho,
                "embedding": embeddings[i] if embeddings else None,
            }
            for i, trecho in enumerate(chunks)
        ]
        # insere em lotes para não estourar o payload
        for i in range(0, len(registros), 50):
            cliente.table("chunks_referencia").insert(registros[i : i + 50]).execute()
        return len(chunks)
    except Exception as exc:  # noqa: BLE001
        raise ErroRAG(f"Falha ao gravar na base de conhecimento: {exc}") from exc


def listar_referencias() -> list[dict]:
    try:
        return (
            db._cliente()  # noqa: SLF001
            .table("documentos_referencia")
            .select("id, titulo, categoria, nome_arquivo, total_chunks, criado_em")
            .order("criado_em", desc=True)
            .execute()
        ).data or []
    except Exception as exc:  # noqa: BLE001
        raise ErroRAG(f"Falha ao listar a base de conhecimento: {exc}") from exc


def excluir_referencia(documento_id: str) -> None:
    try:
        # chunks caem em cascata (on delete cascade)
        db._cliente().table("documentos_referencia").delete().eq(  # noqa: SLF001
            "id", documento_id
        ).execute()
    except Exception as exc:  # noqa: BLE001
        raise ErroRAG(f"Falha ao excluir referência: {exc}") from exc


# ---------------------------------------------------------------------------
# Recuperação (busca) e montagem do bloco de contexto para o prompt
# ---------------------------------------------------------------------------
def _executar_rpc(funcao: str, params: dict) -> list[dict]:
    cliente = db._cliente()  # noqa: SLF001
    try:
        # Fase 2: busca restrita ao tenant do contexto da sessão. Antes da
        # migração 0007 a função não tem o parâmetro `tenant` — o PostgREST
        # devolve erro citando-o e a chamada é repetida na forma antiga
        # (comportamento global de tenant único, idêntico ao de hoje).
        try:
            resposta = cliente.rpc(
                funcao, {**params, "tenant": db.tenant_atual()}
            ).execute()
        except Exception as exc:  # noqa: BLE001
            if "tenant" not in str(exc).lower():
                raise
            resposta = cliente.rpc(funcao, params).execute()
        return resposta.data or []
    except Exception as exc:  # noqa: BLE001
        raise ErroRAG(f"Falha na busca da base de conhecimento: {exc}") from exc


def buscar_referencias(consulta: str, qtd: int = RAG_TOP_K) -> list[dict]:
    """Top-k trechos relevantes: vetorial se possível, senão textual."""
    if not db.disponivel():
        return []
    embedding = _gerar_embeddings([consulta], para_consulta=True)
    if embedding:
        return _executar_rpc("buscar_chunks_vetorial",
                             {"query_embedding": embedding[0], "qtd": qtd})
    return _executar_rpc("buscar_chunks_textual",
                         {"consulta": consulta, "qtd": qtd})


# ---------------------------------------------------------------------------
# Seleção de temas e piso de relevância
# ---------------------------------------------------------------------------
def _gatilhos(dados: dict) -> set[str]:
    """
    Características ESTRUTURADAS da contratação que habilitam temas.
    Reaproveita a derivação canônica de `fatos.py` — o tema não é
    escolhido por uma palavra solta no objeto.
    """
    from . import fatos

    execucao = str(dados.get("modelo_execucao") or "")
    categoria, _ = fatos.categoria_do_objeto(dados)
    campos = fatos._texto_dos_campos(dados)  # noqa: SLF001
    gatilhos = set()
    if execucao.startswith("Sistema de Registro de Preços"):
        gatilhos.add("srp")
    # Proteção de dados NÃO decorre de "ser de TI": um monitor não trata
    # dado pessoal. Dispara para software/SaaS/hospedagem — que operam
    # dados da Administração — ou quando o próprio processo exige o tema.
    if categoria == "TI_SOFTWARE" or _TERMOS_DADOS_NO_PROCESSO.search(
            " ".join(campos.values())):
        gatilhos.add("dados")
    if fatos._tem_termo(campos, fatos._TERMOS_GARANTIA):  # noqa: SLF001
        gatilhos.add("garantia")
    return gatilhos


def temas_para(dados: dict, doc_key: str,
               limite: int = MAX_TEMAS) -> list[str]:
    """
    Lista CONTROLADA de temas: o NÚCLEO do documento (sempre consultado,
    nunca disputa vaga) seguido dos COMPLEMENTARES — primeiro os que o
    objeto/modelagem acionam, depois os demais, se houver folga.
    Tema condicional sem gatilho não é consultado.
    """
    gatilhos = _gatilhos(dados)
    nucleo = [c for c in TEMAS_NUCLEO.get(doc_key, ())][:MAX_TEMAS_NUCLEO]
    acionados, folga = [], []
    for chave in TEMAS_COMPLEMENTARES.get(doc_key, ()):
        if chave in nucleo:
            continue
        condicao = TEMAS_JURIDICOS[chave][2]
        if condicao is None:
            folga.append(chave)
        elif condicao in gatilhos:
            acionados.append(chave)
    complementares = (acionados + folga)[:MAX_TEMAS_COMPLEMENTARES]
    return (nucleo + complementares)[:limite]


def temas_prioritarios(dados: dict, doc_key: str) -> list[str]:
    """Temas que têm direito a uma vaga reservada no bloco (núcleo +
    complementares acionados pelo objeto)."""
    gatilhos = _gatilhos(dados)
    selecionados = temas_para(dados, doc_key)
    return [c for c in selecionados
            if c in TEMAS_NUCLEO.get(doc_key, ())
            or TEMAS_JURIDICOS[c][2] in gatilhos]


def piso_de_relevancia(modo: str) -> float:
    """Piso configurável (config_app); padrão conservador por modo."""
    padrao = PISO_VETORIAL_PADRAO if modo == "vetorial" else PISO_TEXTUAL_PADRAO
    bruto = db.obter_config(f"rag_piso_{modo}") if db.disponivel() else ""
    try:
        return float(str(bruto).replace(",", ".")) if bruto else padrao
    except ValueError:
        return padrao


def _identidade(trecho: dict) -> tuple:
    """Chave de deduplicação: id do chunk ou (documento, ordem)/conteúdo."""
    if trecho.get("id"):
        return ("id", trecho["id"])
    if trecho.get("documento_id") is not None and trecho.get("ordem") is not None:
        return ("pos", trecho["documento_id"], trecho["ordem"])
    return ("txt", (trecho.get("conteudo") or "")[:200])


def _score(trecho: dict) -> float:
    try:
        return float(trecho.get("similaridade") or 0.0)
    except (TypeError, ValueError):
        return 0.0


def _prioridade_fonte(trecho: dict) -> int:
    """Norma > entendimento/acórdão > modelo/processo anterior."""
    categoria = (trecho.get("categoria") or "").lower()
    if categoria == "lei":
        return 3
    if categoria in NORMATIVAS:
        return 2
    if categoria in MOLDES:
        return 0
    return 1


def recuperar(dados: dict, doc_key: str) -> dict:
    """
    Recuperação temática com metadados, piso de relevância, deduplicação
    e teto de trechos. Devolve {"referencias": [...], "consultas": [...],
    "modo": ..., "piso": ...} — as `consultas` são a base do trace.
    Nunca levanta: sem banco/base, devolve estrutura vazia.
    """
    vazio = {"referencias": [], "consultas": [], "modo": "", "piso": None,
             "descartados": 0}
    if not db.disponivel():
        return vazio

    temas = temas_para(dados, doc_key)
    consultas = [{"tema": "geral", "rotulo": "Contexto geral da contratação",
                  "texto": montar_consulta(dados, doc_key)}]
    for chave in temas:
        rotulo, termos, _ = TEMAS_JURIDICOS[chave]
        consultas.append({
            "tema": chave, "rotulo": rotulo,
            "texto": f"{termos} {dados.get('objeto') or ''}".strip()[:500],
        })

    # UMA chamada de embeddings para todas as consultas (custo controlado)
    embeddings = _gerar_embeddings([c["texto"] for c in consultas],
                                   para_consulta=True)
    modo = "vetorial" if embeddings else "textual"
    piso = piso_de_relevancia(modo)

    selecionados: dict[tuple, dict] = {}
    descartados = 0
    for i, consulta in enumerate(consultas):
        qtd = RAG_TOP_K if consulta["tema"] == "geral" else TOP_K_TEMA
        try:
            if embeddings:
                brutos = _executar_rpc(
                    "buscar_chunks_vetorial",
                    {"query_embedding": embeddings[i], "qtd": qtd})
            else:
                brutos = _executar_rpc(
                    "buscar_chunks_textual",
                    {"consulta": consulta["texto"], "qtd": qtd})
        except ErroRAG:
            raise
        consulta["recuperados"] = len(brutos)
        for bruto in brutos:
            score = _score(bruto)
            if score < piso:
                descartados += 1
                continue
            chave = _identidade(bruto)
            existente = selecionados.get(chave)
            if existente is None:
                selecionados[chave] = {
                    **bruto, "tema": consulta["tema"],
                    "tema_rotulo": consulta["rotulo"], "score": score,
                }
            elif score > existente["score"]:
                # mesma fonte recuperada por dois temas: fica o tema em
                # que ela é mais relevante (nunca duplica no prompt)
                existente.update({"tema": consulta["tema"],
                                  "tema_rotulo": consulta["rotulo"],
                                  "score": score})

    ranking = sorted(selecionados.values(),
                     key=lambda t: (_prioridade_fonte(t), t["score"]),
                     reverse=True)
    referencias = _selecionar_com_reserva(
        ranking, temas_prioritarios(dados, doc_key))
    return {"referencias": referencias, "consultas": consultas,
            "modo": modo, "piso": piso, "descartados": descartados}


def _selecionar_com_reserva(ranking: list[dict],
                            prioritarios: list[str]) -> list[dict]:
    """
    Reserva UMA vaga para a melhor evidência de cada tema prioritário e
    só então preenche o restante pelo ranking global.

    Sem isso, um tema com muitos trechos fortes (ex.: requisitos) ocupa
    todas as vagas e o tema consultado ao lado (ex.: sanções) chega ao
    prompt sem nenhuma evidência — a recuperação temática teria sido
    inútil. Nada é duplicado: a reserva usa a mesma lista.
    """
    escolhidos: list[dict] = []
    vistos: set[int] = set()
    for tema in prioritarios:
        melhor = next((t for t in ranking
                       if t.get("tema") == tema and id(t) not in vistos), None)
        if melhor is not None and len(escolhidos) < MAX_CHUNKS_PROMPT:
            escolhidos.append(melhor)
            vistos.add(id(melhor))
    for trecho in ranking:
        if len(escolhidos) >= MAX_CHUNKS_PROMPT:
            break
        if id(trecho) not in vistos:
            escolhidos.append(trecho)
            vistos.add(id(trecho))
    # ordem final de exibição: hierarquia da fonte e relevância
    return sorted(escolhidos,
                  key=lambda t: (_prioridade_fonte(t), t["score"]),
                  reverse=True)


def montar_consulta(dados: dict, doc_key: str) -> str:
    """Texto de busca combinando objeto, justificativa e o tipo de documento."""
    nomes = {
        "dfd": "documento de formalização da demanda",
        "etp": "estudo técnico preliminar",
        "tr": "termo de referência",
        "edital": "edital de licitação registro de preços",
    }
    partes = [
        nomes[doc_key],
        dados.get("objeto") or "",
        dados.get("justificativa") or "",
        dados.get("modelo_execucao") or "",
    ]
    return " ".join(p for p in partes if p)[:1500]


# Regra de citação (P1): o número do dispositivo só entra no documento
# quando sustentado por trecho recuperado ou pelo mapa canônico do system
# prompt. Sem lastro, cita-se a norma sem o artigo — nunca se inventa.
REGRA_DE_CITACAO = (
    "REGRA DE CITAÇÃO (obrigatória): só escreva o NÚMERO de um dispositivo "
    "(artigo, inciso, parágrafo, súmula, acórdão, decreto, instrução "
    "normativa) quando ele estiver (a) presente em um dos trechos "
    "recuperados abaixo ou (b) no mapa canônico da Lei nº 14.133/2021 "
    "constante das suas instruções. Se você sabe a norma aplicável mas não "
    "tem lastro para o dispositivo específico, cite apenas a norma — "
    "'nos termos da Lei nº 14.133/2021' é preferível a um artigo errado. "
    "É PROIBIDO deduzir número de artigo por memória ou analogia."
)

_HIERARQUIA_FONTES = (
    "HIERARQUIA DAS FONTES (respeite estritamente):\n"
    "1) LEGISLAÇÃO E REGULAMENTO (lei, decreto, regulamento municipal): "
    "fundamenta diretamente a cláusula.\n"
    "2) JURISPRUDÊNCIA E ORIENTAÇÃO DE CONTROLE (acórdãos e "
    "entendimentos de Tribunais de Contas): orienta a INTERPRETAÇÃO da "
    "norma e as boas práticas de controle. NÃO é legislação: cite-a como "
    "orientação ('conforme entendimento do TCU no Acórdão nº …'), nunca "
    "como se fosse o texto da lei, e jamais deduza dela um número de "
    "artigo de lei.\n"
    "3) PROCESSO ATUAL (memorando, formulário, planilha, anexos): é a "
    "única fonte dos FATOS desta contratação.\n"
    "4) PROCESSO ANTERIOR / MODELO PADRÃO: fonte apenas de ESTRUTURA, "
    "ordem dos tópicos, linguagem e cláusulas institucionais recorrentes. "
    "NÃO é prova do direito vigente nem fonte de fato: é PROIBIDO "
    "transportar dele objeto, justificativa, quantitativos, valores, "
    "fornecedores, fiscais/gestores, dotações, unidades, prazos, datas ou "
    "números, e é PROIBIDO citar um dispositivo apenas porque ele "
    "aparecia no documento antigo (a norma pode ter mudado).\n"
    "JURISDIÇÃO: esta é uma contratação MUNICIPAL. Aplicam-se diretamente "
    "a Lei nº 14.133/2021 e a regulamentação do próprio Município, que "
    "tem PRECEDÊNCIA OPERACIONAL quando disciplinar a matéria. Instruções "
    "normativas, decretos e manuais FEDERAIS só valem como referência "
    "técnica — nunca os apresente como norma obrigatória para o Município "
    "sem que o trecho recuperado demonstre a aplicabilidade."
)


_RE_ARTIGO_TRECHO = re.compile(r"\bart(?:igo|s?)?\.?\s*(\d{1,3})\s*[º°]?",
                               re.IGNORECASE)


def dispositivos_do_trecho(texto: str, titulo: str = "") -> list[str]:
    """
    Dispositivos citados no trecho, na forma `norma:artigo`.

    O número isolado não identifica nada: o art. 84 da Lei nº 14.133/2021
    não autoriza um "art. 84" de outra norma. A norma vem do próprio
    trecho ou, na falta dela, do TÍTULO da fonte indexada (é o que
    identifica o documento na base).
    """
    from .normas import identificar_norma

    norma = identificar_norma(texto) or identificar_norma(titulo)
    if not norma:
        return []
    return sorted({f"{norma}:{m.group(1)}"
                   for m in _RE_ARTIGO_TRECHO.finditer(texto or "")},
                  key=lambda d: int(d.split(":")[1]))


def lastro_do_trace(trace: dict | None) -> set[str]:
    """
    Dispositivos (`norma:artigo`) com lastro nas fontes recuperadas.

    Só LEGISLAÇÃO sustenta dispositivo normativo. Um acórdão pode citar o
    art. 40 ao interpretá-lo, mas quem autoriza o documento a invocar o
    art. 40 é a lei — não a ementa que o menciona. Processo anterior e
    modelo nunca fornecem lastro (a norma pode ter mudado).
    """
    return {d for r in (trace or {}).get("referencias", [])
            if (r.get("categoria") or "").lower() in LEGISLACAO
            for d in (r.get("dispositivos") or [])}


def montar_contexto(dados: dict, doc_key: str) -> dict:
    """
    {"bloco": texto para o prompt, "trace": rastro auditável}.
    O bloco agrupa os trechos POR TEMA e expõe fonte, categoria e
    relevância — o modelo enxerga a evidência de cada matéria, não uma
    pilha indistinta. Nunca levanta exceção: RAG é enriquecimento.
    """
    trace = {"consultas": [], "referencias": [], "modo": "", "piso": None}
    try:
        resultado = recuperar(dados, doc_key)
    except ErroRAG as erro:
        st.warning(str(erro))
        trace["erro"] = str(erro)[:200]
        return {"bloco": "", "trace": trace}

    referencias = resultado["referencias"]
    trace = {
        "modo": resultado["modo"],
        "piso": resultado["piso"],
        "descartados_por_piso": resultado["descartados"],
        "consultas": [
            {"tema": c["tema"], "consulta": c["texto"][:200],
             "recuperados": c.get("recuperados", 0)}
            for c in resultado["consultas"]
        ],
        # rastro da FONTE, não do conteúdo: título, categoria, score e
        # posição bastam para responder "por que citou este artigo?"
        "referencias": [
            {"tema": r.get("tema"), "titulo": r.get("titulo"),
             "categoria": r.get("categoria"),
             "score": round(float(r.get("score") or 0), 4),
             "documento_id": r.get("documento_id"), "ordem": r.get("ordem"),
             # dispositivos EXPRESSOS no trecho (norma:artigo): é o
             # lastro que autoriza o documento a citar o dispositivo
             "dispositivos": dispositivos_do_trecho(r.get("conteudo"),
                                                    r.get("titulo") or ""),
             "trecho": (r.get("conteudo") or "")[:160]}
            for r in referencias
        ],
    }
    if not referencias:
        return {"bloco": "", "trace": trace}

    linhas = [
        "\n=== REFERÊNCIAS DA BASE DE CONHECIMENTO (recuperadas por tema) ===",
        _HIERARQUIA_FONTES,
        REGRA_DE_CITACAO,
        "Onde faltar dado do processo atual, use [PREENCHER] — nunca "
        "preencha com dado de outro processo.",
    ]
    por_tema: dict[str, list[dict]] = {}
    for referencia in referencias:
        por_tema.setdefault(referencia.get("tema_rotulo") or "Geral",
                            []).append(referencia)
    for tema, itens in por_tema.items():
        linhas.append(f"\n### TEMA: {tema}")
        for i, t in enumerate(itens, start=1):
            categoria = (t.get("categoria") or "").lower()
            rotulo = CATEGORIAS.get(categoria, categoria or "Outro")
            papel = _PAPEL_DA_FONTE.get(
                categoria, "fonte não classificada — trate com cautela e "
                           "não a use para fundamentar dispositivo")
            linhas.append(
                f"[{i}] Fonte: {t.get('titulo', '(sem título)')} "
                f"| Tipo: {rotulo} ({papel}) "
                f"| Relevância: {t.get('score', 0):.3f}")
            linhas.append("Trecho recuperado: "
                          + (t.get("conteudo") or "").strip())
    return {"bloco": "\n".join(linhas), "trace": trace}


def montar_bloco_referencias(dados: dict, doc_key: str) -> str:
    """Compatibilidade: apenas o bloco de texto do contexto recuperado."""
    return montar_contexto(dados, doc_key)["bloco"]
