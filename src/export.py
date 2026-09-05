"""
Exportação dos documentos aprovados para .docx e .pdf.

Os textos gerados pela IA vêm em Markdown simples (títulos #/##/###,
listas e negrito **texto**). Este módulo converte esse Markdown em um
DOCX ESTRUTURADO com os estilos institucionais dos documentos aprovados
(Times New Roman 12, espaçamento 1,5, 6 pt após parágrafo, texto
justificado, cláusulas numeradas em negrito, controle de linhas órfãs e
título preso ao conteúdo). O PDF é obtido preferencialmente CONVERTENDO
esse DOCX via LibreOffice — garantindo que DOCX e PDF tenham o mesmo
conteúdo e a mesma formatação; sem LibreOffice no ambiente, cai para o
renderizador fpdf2 (fonte Times), com aviso via motor_pdf().
"""

import io
import re
import shutil
import zipfile
from datetime import date

from .config import DOCUMENTOS, DOCUMENTOS_EXPORTAVEIS, exportaveis_do_processo


def _ordem_de_exportacao(documentos: dict[str, str],
                         dados: dict | None) -> list[str]:
    """
    Quais documentos entram no arquivo, e em que ordem.

    Com o formulário do processo em mãos, quem decide é
    `config.exportaveis_do_processo` — e um processo sem Sistema de
    Registro de Preços não exporta ARP nem que exista uma chave 'arp'
    residual de uma modelagem anterior. Sem o formulário (minuta de
    trabalho, bundle avulso), exporta-se o que veio: aqui não há como
    saber se a Ata cabe, e inventar a resposta seria pior.
    """
    if dados is None:
        return [k for k in DOCUMENTOS_EXPORTAVEIS if k in documentos]
    return exportaveis_do_processo(dados, documentos)

# Padrão institucional (medido nos documentos manuais aprovados)
FONTE_CORPO = "Times New Roman"
TAMANHO_CORPO = 12          # pt
ESPACO_LINHAS = 1.5
ESPACO_DEPOIS = 6           # pt após parágrafos
MARGEM_SUP_CM = 2.5
MARGEM_INF_CM = 2.5
MARGEM_ESQ_CM = 2.0
MARGEM_DIR_CM = 2.0

# ---------------------------------------------------------------------------
# Utilitários de parsing do Markdown simplificado
# ---------------------------------------------------------------------------
_RE_NEGRITO = re.compile(r"\*\*(.+?)\*\*")


def _classificar_linha(linha: str) -> tuple[str, str]:
    """Retorna (tipo, conteúdo) para cada linha do Markdown."""
    txt = linha.rstrip()
    if not txt.strip():
        return "vazio", ""
    if txt.startswith("### "):
        return "h3", txt[4:].strip()
    if txt.startswith("## "):
        return "h2", txt[3:].strip()
    if txt.startswith("# "):
        return "h1", txt[2:].strip()
    if re.match(r"^\s*[-*]\s+", txt):
        return "item", re.sub(r"^\s*[-*]\s+", "", txt)
    if txt.strip().startswith("|"):
        return "tabela", txt.strip()
    return "par", txt


def _limpar_inline(texto: str) -> str:
    """Remove marcações inline (negrito/itálico) para saídas sem rich text."""
    texto = _RE_NEGRITO.sub(r"\1", texto)
    return texto.replace("*", "")


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def _definir_fonte(estilo, nome: str, tamanho_pt: float, negrito: bool = False):
    """Aplica a fonte também em rFonts hAnsi/cs (Word ignora só o ascii)."""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    estilo.font.name = nome
    estilo.font.size = Pt(tamanho_pt)
    estilo.font.bold = negrito
    rpr = estilo.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), nome)


def _novo_estilo(doc, nome: str, base: str = "Normal"):
    from docx.enum.style import WD_STYLE_TYPE

    try:
        estilo = doc.styles[nome]
    except KeyError:
        estilo = doc.styles.add_style(nome, WD_STYLE_TYPE.PARAGRAPH)
        if base:
            estilo.base_style = doc.styles[base]
    return estilo


def _docx_novo():
    """
    Documento A4 com os ESTILOS INSTITUCIONAIS centralizados (nada de
    formatação manual parágrafo a parágrafo):
      GovDocs Corpo / Titulo / Clausula / Item 1..3 / Tabela / Nota / Assinatura.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.shared import Cm, Pt

    doc = Document()
    secao = doc.sections[0]
    secao.page_width, secao.page_height = Cm(21.0), Cm(29.7)  # A4
    secao.top_margin, secao.bottom_margin = Cm(MARGEM_SUP_CM), Cm(MARGEM_INF_CM)
    secao.left_margin, secao.right_margin = Cm(MARGEM_ESQ_CM), Cm(MARGEM_DIR_CM)

    def _paragrafo(estilo, *, alinhamento=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   depois=ESPACO_DEPOIS, linhas=ESPACO_LINHAS, recuo_cm=0.0,
                   manter_com_proximo=False):
        pf = estilo.paragraph_format
        pf.alignment = alinhamento
        pf.space_after = Pt(depois)
        pf.space_before = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = linhas
        pf.widow_control = True            # sem linhas órfãs/viúvas
        pf.keep_with_next = manter_com_proximo
        if recuo_cm:
            pf.left_indent = Cm(recuo_cm)

    # Normal = corpo (herdado por tudo)
    normal = doc.styles["Normal"]
    _definir_fonte(normal, FONTE_CORPO, TAMANHO_CORPO)
    _paragrafo(normal)

    corpo = _novo_estilo(doc, "GovDocs Corpo")
    _definir_fonte(corpo, FONTE_CORPO, TAMANHO_CORPO)
    _paragrafo(corpo)

    titulo = _novo_estilo(doc, "GovDocs Titulo")
    _definir_fonte(titulo, FONTE_CORPO, 14, negrito=True)
    _paragrafo(titulo, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, depois=12,
               manter_com_proximo=True)

    clausula = _novo_estilo(doc, "GovDocs Clausula")
    _definir_fonte(clausula, FONTE_CORPO, TAMANHO_CORPO, negrito=True)
    _paragrafo(clausula, alinhamento=WD_ALIGN_PARAGRAPH.LEFT, depois=6,
               manter_com_proximo=True)  # título nunca separa do 1º parágrafo

    for nome, recuo in (("GovDocs Item 1", 0.75), ("GovDocs Item 2", 1.5),
                        ("GovDocs Item 3", 2.25)):
        item = _novo_estilo(doc, nome)
        _definir_fonte(item, FONTE_CORPO, TAMANHO_CORPO)
        _paragrafo(item, recuo_cm=recuo)

    nota = _novo_estilo(doc, "GovDocs Nota")
    _definir_fonte(nota, FONTE_CORPO, 10)
    _paragrafo(nota, depois=4, linhas=1.0)

    assin = _novo_estilo(doc, "GovDocs Assinatura")
    _definir_fonte(assin, FONTE_CORPO, TAMANHO_CORPO)
    _paragrafo(assin, alinhamento=WD_ALIGN_PARAGRAPH.CENTER, depois=0,
               manter_com_proximo=True)  # bloco de assinatura não divide
    return doc


# Links Markdown: [texto](url) — usados para compactar URLs (fonte de preço)
_RE_LINK = re.compile(r"\[([^\]]+)\]\((https?://[^\s)]+)\)")


def _segmentos_bold(texto: str) -> list[dict]:
    segs, pos = [], 0
    for m in _RE_NEGRITO.finditer(texto):
        if m.start() > pos:
            segs.append({"text": texto[pos : m.start()], "bold": False, "url": None})
        segs.append({"text": m.group(1), "bold": True, "url": None})
        pos = m.end()
    if pos < len(texto):
        segs.append({"text": texto[pos:], "bold": False, "url": None})
    return segs


def _segmentos_ricos(texto: str) -> list[dict]:
    """Divide o texto em trechos: negrito (**), links [t](url) e texto simples."""
    segs, pos = [], 0
    for m in _RE_LINK.finditer(texto):
        if m.start() > pos:
            segs.extend(_segmentos_bold(texto[pos : m.start()]))
        segs.append({"text": m.group(1), "bold": False, "url": m.group(2)})
        pos = m.end()
    if pos < len(texto):
        segs.extend(_segmentos_bold(texto[pos:]))
    return segs or [{"text": texto, "bold": False, "url": None}]


def _docx_hyperlink(par, url: str, texto: str) -> None:
    """Insere um hyperlink clicável (azul, sublinhado) no parágrafo."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    r_id = par.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    cor = OxmlElement("w:color"); cor.set(qn("w:val"), "1B4F8A")
    sub = OxmlElement("w:u"); sub.set(qn("w:val"), "single")
    rpr.append(cor); rpr.append(sub); run.append(rpr)
    t = OxmlElement("w:t"); t.text = texto; run.append(t)
    link.append(run)
    par._p.append(link)


def _docx_runs_ricos(par, texto: str) -> None:
    """Preenche o parágrafo com runs de negrito e hyperlinks compactos."""
    for seg in _segmentos_ricos(texto):
        if seg["url"]:
            _docx_hyperlink(par, seg["url"], seg["text"])
        elif seg["text"]:
            run = par.add_run(seg["text"])
            if seg["bold"]:
                # só marca quando positivo: run.bold=False anularia o negrito
                # herdado do estilo (ex.: títulos de cláusula)
                run.bold = True


def _docx_paragrafo_com_negrito(doc, texto: str, estilo: str | None = None):
    """Adiciona parágrafo com negrito (**) e links clicáveis [t](url)."""
    par = doc.add_paragraph(style=estilo)
    _docx_runs_ricos(par, texto)
    return par


# parágrafos que começam com numeração hierárquica: 1.1. / 1.1.1. / 1.1.1.1.
_RE_NIVEL = re.compile(r"^\s*\d{1,2}(\.\d{1,2}){1,3}\.?\s")


def _estilo_do_paragrafo(texto: str) -> str:
    """Estilo institucional conforme a profundidade da numeração do texto."""
    if "____" in texto:
        return "GovDocs Assinatura"
    m = _RE_NIVEL.match(texto)
    if not m:
        return "GovDocs Corpo"
    profundidade = m.group(0).count(".")  # 1.1.=2  1.1.1.=3  1.1.1.1.=4
    return {2: "GovDocs Item 1", 3: "GovDocs Item 2"}.get(profundidade,
                                                          "GovDocs Item 3")


# Uma linha de tabela com até este total de caracteres cabe com folga em
# uma página; acima disso (descrições longas de planilha) ela PRECISA poder
# dividir entre páginas — senão cada item ocupa uma página inteira.
_LIMITE_LINHA_SEM_QUEBRA = 250

# Linha separadora de cabeçalho Markdown: |---|---| (com ou sem ':')
_RE_SEPARADOR_TABELA = re.compile(r"^\|?[\s:|-]+\|?$")

# Divisão de células que RESPEITA o escape `\|`.
#
# O `split("|")` ingênuo transformava dado em ESTRUTURA: uma descrição de
# item contendo "CANETA | 999999,00" — vinda de fonte externa ou digitada
# na planilha do processo — acrescentava colunas à tabela do documento
# oficial e deslocava os valores para a coluna errada. Medido: uma tabela
# de 13 colunas virava 14, com o número forjado caindo sob "Descrição".
#
# Quem escreve a tabela escapa a barra (`\|`); aqui ela é reconhecida
# como conteúdo e volta a ser uma barra comum na célula.
_RE_PIPE_ESTRUTURAL = re.compile(r"(?<!\\)\|")


def _celulas_da_linha(linha: str) -> list[str]:
    r"""Células de uma linha de tabela Markdown, com `\|` como conteúdo."""
    corpo = linha.strip()
    if corpo.startswith("|"):
        corpo = corpo[1:]
    # A barra final só é delimitador se não estiver escapada.
    if corpo.endswith("|") and not corpo.endswith("\\|"):
        corpo = corpo[:-1]
    return [celula.strip().replace("\\|", "|")
            for celula in _RE_PIPE_ESTRUTURAL.split(corpo)]


def _tem_cabecalho(tabela_buffer: list[str]) -> bool:
    """
    A tabela Markdown só tem cabeçalho REAL quando a 2ª linha é a
    separadora (|---|). Sem isso, a 1ª linha é DADO — promovê-la a
    cabeçalho (negrito + repetição em toda página) foi a causa do
    "primeiro item repetido em todas as páginas" nos documentos reais.
    """
    return len(tabela_buffer) >= 2 and bool(
        _RE_SEPARADOR_TABELA.match(tabela_buffer[1].strip()))


# Largura útil da página (A4 21 cm menos as margens laterais)
_LARGURA_UTIL_CM = 21.0 - MARGEM_ESQ_CM - MARGEM_DIR_CM
# Piso e teto por coluna: abaixo do piso o texto quebra a cada poucos
# caracteres; acima do teto uma coluna sozinha come a linha inteira.
_COL_MIN_CM = 1.4
_COL_MAX_CM = 9.5
# Corpo da tabela no DOCX (ver _docx_formatar_tabela)
_TABELA_PT = 10
# Margem interna da célula, explícita: o padrão do Word é 0,19 cm de cada
# lado, e a largura DECLARADA de uma coluna não é a largura utilizável.
# Era essa diferença que fazia "572704" caber na conta e não caber na
# página. Reduzida ao mínimo legível e descontada do cálculo.
_PADDING_CELULA_CM = 0.08


def _texto_renderizado(celula: str) -> str:
    """Texto como sai na página: link Markdown vale pelo seu rótulo."""
    return _limpar_inline(_RE_LINK.sub(r"\1", celula or "")).strip()


# Substitutos latin-1 para os caracteres que MEDEM parecido e que
# aparecem o tempo todo em texto colado do Word ou extraído de PDF. A
# medida sai correta; o documento continua com o caractere original,
# porque quem o escreve é o LibreOffice, não o medidor.
_EQUIVALENTE_PARA_MEDIDA = {
    "—": "--",   # travessão
    "–": "-",    # meia-risca
    "‘": "'", "’": "'",
    "“": '"', "”": '"',
    "…": "...",
    " ": " ",
    "−": "-",    # sinal de menos
    "•": "*",
}

# Largura de reserva, em pontos, para um caractere que nem o mapa acima
# resolve. É o "m" do Times a 10 pt, arredondado para cima: superestimar
# a largura faz a coluna ficar um pouco larga; subestimar faz o conteúdo
# estourar a página, que é o defeito caro.
_LARGURA_DESCONHECIDO_PT = 10.0


def _para_medida(texto: str) -> str:
    """Troca o que o Times latin-1 não encoda por equivalente de largura."""
    saida = []
    for caractere in texto:
        equivalente = _EQUIVALENTE_PARA_MEDIDA.get(caractere)
        if equivalente is not None:
            saida.append(equivalente)
            continue
        try:
            caractere.encode("latin-1")
        except UnicodeEncodeError:
            saida.append(None)          # marcado para largura de reserva
        else:
            saida.append(caractere)
    return saida


def _largura_de_texto_cm(texto: str, pt: float = _TABELA_PT,
                         negrito: bool = False) -> float:
    """
    Largura real de um texto em cm, pelas métricas da fonte Times.

    Medida, não estimada: o fpdf2 já embarca as métricas do Times e é
    dependência do projeto. Uma constante de "largura média de caractere"
    erraria justamente onde importa — dígitos e maiúsculas são mais
    largos que a média, e é neles que a coluna estoura.

    O medidor usa Times em latin-1, e travessão, meia-risca e aspas
    curvas — que vêm em toda descrição colada do Word ou extraída de PDF
    — estão fora dessa faixa. Antes, um único desses caracteres numa
    célula levantava `FPDFUnicodeEncodingException` e derrubava a
    geração INTEIRA do DOCX e do PDF. Medir é um cálculo auxiliar: ele
    não pode ser o que impede o documento de existir.

    A saída é conservadora por construção — caractere desconhecido conta
    como um "m". Superestimar deixa a coluna um pouco larga; subestimar
    estoura a página, que é o defeito que este módulo existe para evitar.
    """
    from fpdf import FPDF

    global _medidor
    if _medidor is None:
        _medidor = FPDF(unit="pt")
        _medidor.add_page()
    _medidor.set_font("Times", "B" if negrito else "", pt)

    pedacos = _para_medida(texto or "")
    mensuraveis = "".join(c for c in pedacos if c is not None)
    desconhecidos = sum(1 for c in pedacos if c is None)
    largura_pt = _medidor.get_string_width(mensuraveis)
    largura_pt += desconhecidos * _LARGURA_DESCONHECIDO_PT * (pt / 10.0)
    return largura_pt / 72 * 2.54


_medidor = None


def _piso_das_colunas_cm(linhas: list[list[str]]) -> list[float]:
    """
    Largura mínima de cada coluna: a do seu maior TOKEN INDIVISÍVEL.

    Uma palavra sem espaço não quebra em lugar nenhum — ou a coluna a
    comporta, ou o conversor a parte no meio. Foi o que aconteceu no PDF
    real: o código '572704' saiu como '57270' + '4' e o cabeçalho
    'Quantidade' como 'Quanti' + 'dade', porque o piso era um número fixo
    (1,4 cm) que não olhava para o texto.
    """
    if not linhas:
        return []
    pisos = []
    for j in range(len(linhas[0])):
        maior = 0.0
        for i, linha in enumerate(linhas):
            # a 1ª linha é medida em NEGRITO mesmo quando não é cabeçalho:
            # negrito é mais largo, e errar para o lado largo só custa
            # alguns milímetros — errar para o estreito parte a palavra
            for token in _texto_renderizado(linha[j]).split():
                maior = max(maior,
                            _largura_de_texto_cm(token, negrito=(i == 0)))
        pisos.append(max(maior, 0.0) + 2 * _PADDING_CELULA_CM)
    return pisos


def _acomodar_na_largura_util(larguras: list[float],
                              pisos: list[float]) -> list[float]:
    """
    Ajusta as larguras à largura útil sem espremer coluna abaixo do piso.

    O excedente sai das colunas FOLGADAS, proporcionalmente à folga que
    cada uma tem sobre o próprio piso — a Descrição cede, o Código não.
    Se nem os pisos couberem (tabela larga demais para A4), tudo é
    reduzido junto: aí não há solução sem quebrar texto, e reduzir de
    forma uniforme ao menos não escolhe uma vítima.
    """
    if not larguras:
        return larguras
    larguras = [max(l, p) for l, p in zip(larguras, pisos)]
    excesso = sum(larguras) - _LARGURA_UTIL_CM
    if excesso <= 0:
        return larguras
    folga_total = sum(l - p for l, p in zip(larguras, pisos))
    if folga_total <= 0:
        fator = _LARGURA_UTIL_CM / sum(larguras)
        return [l * fator for l in larguras]
    corte = min(excesso, folga_total)
    return [l - (l - p) / folga_total * corte
            for l, p in zip(larguras, pisos)]


def _remover_colunas_vazias(linhas: list[list[str]]) -> list[list[str]]:
    """
    Descarta colunas sem nenhum conteúdo.

    O Markdown de tabela costuma gerar uma coluna extra vazia (linha
    terminada em '|'), e no PDF auditado ela aparecia como uma faixa em
    branco em todas as ~150 páginas de tabela, roubando largura das
    colunas que tinham texto.
    """
    if not linhas:
        return linhas
    manter = [j for j in range(len(linhas[0]))
              if any((linha[j] or "").strip() for linha in linhas)]
    if not manter or len(manter) == len(linhas[0]):
        return linhas
    return [[linha[j] for j in manter] for linha in linhas]


def _pesos_das_colunas(linhas: list[list[str]]) -> tuple[float, ...]:
    """
    Peso relativo de cada coluna: raiz da maior célula RENDERIZADA.

    A raiz amortece — uma descrição de 200 caracteres fica larga sem
    espremer as demais a zero. Os pesos são limitados pela mesma razão
    mínimo/máximo usada no DOCX, para que os dois formatos saiam com a
    mesma proporção.
    """
    if not linhas:
        return ()
    n_cols = len(linhas[0])
    pesos = []
    for j in range(n_cols):
        maior = max((len(_texto_renderizado(linha[j])) for linha in linhas),
                    default=1)
        pesos.append(max(maior, 1) ** 0.5)
    total = sum(pesos) or 1
    # normaliza para a largura útil e aplica teto em centímetros
    cm = [min(_LARGURA_UTIL_CM * p / total, _COL_MAX_CM) for p in pesos]
    # o piso é o do MAIOR TOKEN de cada coluna — nunca um número fixo
    pisos = [max(p, _COL_MIN_CM) for p in _piso_das_colunas_cm(linhas)]
    return tuple(round(c, 2) for c in _acomodar_na_largura_util(cm, pisos))


def _docx_larguras_proporcionais(tabela, linhas: list[list[str]]) -> None:
    """
    Distribui a largura entre as colunas conforme o texto que cada uma
    carrega, em vez de reparti-la em partes iguais.

    Sem isto, a Descrição do item (às vezes 200 caracteres) recebia a
    mesma largura de 'Unidade' e o conversor quebrava palavra a cada
    poucos caracteres ("ESPECIFICA ÇÃO", "PASTA SANFONAD A"): no PDF
    auditado eram 455 fragmentos de 1 a 3 letras, e o texto extraído
    ficava sem nem poder ser pesquisado.
    """
    from docx.shared import Cm

    n_cols = len(linhas[0])
    if not n_cols:
        return
    # peso = maior célula da coluna, amortecida (raiz) para que uma
    # descrição muito longa não zere as demais. Mede o texto RENDERIZADO:
    # "[link](https://…60 caracteres…)" ocupa 4 caracteres na página.
    pesos = []
    for j in range(n_cols):
        maior = max((len(_texto_renderizado(linha[j])) for linha in linhas),
                    default=1)
        pesos.append(max(maior, 1) ** 0.5)

    total = sum(pesos) or 1
    larguras = [min(_LARGURA_UTIL_CM * p / total, _COL_MAX_CM) for p in pesos]
    # Piso pelo maior TOKEN de cada coluna, e não por um número fixo: era
    # o piso fixo de 1,4 cm que partia '572704' em '57270' + '4'. O que
    # exceder a largura útil sai das colunas com folga sobre o próprio
    # piso — a Descrição cede, o Código não.
    pisos = [max(p, _COL_MIN_CM) for p in _piso_das_colunas_cm(linhas)]
    larguras = _acomodar_na_largura_util(larguras, pisos)

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tabela.autofit = False
    for j, largura in enumerate(larguras):
        for linha in tabela.rows:
            linha.cells[j].width = Cm(round(largura, 2))

    # A GRADE da tabela (w:tblGrid) é o que o conversor usa para repartir
    # a largura; sem atualizá-la, as larguras de célula são ignoradas.
    grid = tabela._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for col, largura in zip(grid.findall(qn("w:gridCol")), larguras):
            col.set(qn("w:w"), str(int(Cm(round(largura, 2)).twips)))

    # …e a LARGURA TOTAL (w:tblW) precisa ser explícita: com ela ausente
    # o LibreOffice recalcula a tabela do zero e volta a distribuir as
    # colunas em partes iguais, desfazendo a grade acima.
    #
    # A POSIÇÃO importa: o OOXML fixa a ordem dos filhos de <w:tblPr> e
    # <w:tblW> vem ANTES de <w:tblLayout>. Anexado ao fim, o elemento é
    # descartado silenciosamente pelo conversor — foi o que aconteceu na
    # primeira tentativa desta correção.
    tbl_pr = tabela._tbl.tblPr
    for antigo in tbl_pr.findall(qn("w:tblW")):
        tbl_pr.remove(antigo)
    largura_total = OxmlElement("w:tblW")
    largura_total.set(qn("w:type"), "dxa")
    largura_total.set(qn("w:w"), str(int(Cm(_LARGURA_UTIL_CM).twips)))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is not None:
        layout.addprevious(largura_total)
    else:
        tbl_pr.append(largura_total)


def _docx_formatar_tabela(tabela, com_cabecalho: bool = True) -> None:
    """Cabeçalho repetido por página (quando existe), fonte do padrão e
    quebra de página permitida nas linhas longas."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    # Margem interna EXPLÍCITA. O padrão do Word é 0,19 cm de cada lado, e
    # o cálculo de largura não a enxergava: a coluna do Código recebia
    # 1,4 cm, sobravam 1,02 cm de texto útil, e '572704' (1,06 cm em
    # Times 10) saía partido como '57270' + '4' no PDF convertido.
    tbl_pr = tabela._tbl.tblPr
    for antiga in tbl_pr.findall(qn("w:tblCellMar")):
        tbl_pr.remove(antiga)
    margens = OxmlElement("w:tblCellMar")
    for lado, cm in (("top", 0.03), ("left", _PADDING_CELULA_CM),
                     ("bottom", 0.03), ("right", _PADDING_CELULA_CM)):
        elemento = OxmlElement(f"w:{lado}")
        elemento.set(qn("w:w"), str(int(Cm(cm).twips)))
        elemento.set(qn("w:type"), "dxa")
        margens.append(elemento)
    tbl_pr.append(margens)

    for i, linha in enumerate(tabela.rows):
        tr_pr = linha._tr.get_or_add_trPr()
        eh_cabecalho = com_cabecalho and i == 0
        if eh_cabecalho:  # repete o cabeçalho nas páginas seguintes
            cab = OxmlElement("w:tblHeader")
            cab.set(qn("w:val"), "true")
            tr_pr.append(cab)
        total_chars = sum(len(cel.text) for cel in linha.cells)
        if total_chars <= _LIMITE_LINHA_SEM_QUEBRA:
            sem_quebra = OxmlElement("w:cantSplit")  # linha curta não divide
            tr_pr.append(sem_quebra)
        for cel in linha.cells:
            for par in cel.paragraphs:
                pf = par.paragraph_format
                pf.space_after = Pt(2)
                pf.line_spacing = 1.0
                for run in par.runs:
                    run.font.name = FONTE_CORPO
                    run.font.size = Pt(10)
                    if eh_cabecalho:
                        run.font.bold = True


def _docx_inserir_markdown(doc, texto_md: str) -> None:
    linhas = texto_md.splitlines()
    tabela_buffer: list[str] = []

    def descarregar_tabela():
        if not tabela_buffer:
            return
        com_cabecalho = _tem_cabecalho(tabela_buffer)
        linhas_tab = [
            _celulas_da_linha(ln)
            for ln in tabela_buffer
            if not _RE_SEPARADOR_TABELA.match(ln)  # descarta linha ---|---
        ]
        if linhas_tab:
            n_cols = max(len(l) for l in linhas_tab)
            linhas_tab = [l + [""] * (n_cols - len(l)) for l in linhas_tab]
            linhas_tab = _remover_colunas_vazias(linhas_tab)
            n_cols = len(linhas_tab[0])
            tabela = doc.add_table(rows=len(linhas_tab), cols=n_cols)
            tabela.style = "Table Grid"
            # Preenchimento por LINHA, e não por `tabela.cell(i, j)`.
            #
            # No python-docx, `Table.cell` passa por `_cells`, que
            # reconstrói a lista de TODAS as células da tabela a cada
            # acesso. Preencher célula a célula é, portanto, quadrático
            # no tamanho da tabela: medido com um relatório de 15 itens,
            # `cell()` respondia por 14,4 s de 18,5 s. `row.cells`
            # constrói a lista uma vez por linha.
            for i, (linha_docx, linha) in enumerate(
                    zip(tabela.rows, linhas_tab)):
                celulas_docx = linha_docx.cells
                for j, celula in enumerate(linha[: len(celulas_docx)]):
                    par = celulas_docx[j].paragraphs[0]
                    if i == 0 and com_cabecalho:  # cabeçalho: negrito, sem links
                        par.add_run(_limpar_inline(celula)).bold = True
                    else:
                        _docx_runs_ricos(par, celula)
            _docx_larguras_proporcionais(tabela, linhas_tab)
            _docx_formatar_tabela(tabela, com_cabecalho)
        tabela_buffer.clear()

    for linha in linhas:
        tipo, conteudo = _classificar_linha(linha)
        if tipo == "tabela":
            tabela_buffer.append(conteudo)
            continue
        descarregar_tabela()
        if tipo in ("h1", "h2", "h3"):
            # cláusulas numeradas em negrito, presas ao 1º parágrafo
            _docx_paragrafo_com_negrito(
                doc, _limpar_inline(conteudo), estilo="GovDocs Clausula")
        elif tipo == "item":
            _docx_paragrafo_com_negrito(doc, "•  " + conteudo,
                                        estilo="GovDocs Item 1")
        elif tipo == "par":
            _docx_paragrafo_com_negrito(doc, conteudo,
                                        estilo=_estilo_do_paragrafo(conteudo))
    descarregar_tabela()


def _docx_aplicar_branding(doc, branding: dict | None) -> None:
    """
    Cabeçalho e rodapé institucionais no DOCX. Prioriza IMAGENS capturadas
    do modelo (inseridas na largura do conteúdo, na área de header/footer);
    sem imagem, usa TEXTO. A marca d'água translúcida é aplicada apenas no
    PDF (limitação do formato DOCX).
    """
    if not branding:
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    img_cab = _img_bytes(branding, "cabecalho_img")
    img_rod = _img_bytes(branding, "rodape_img")
    secao = doc.sections[0]
    largura_conteudo = secao.page_width - secao.left_margin - secao.right_margin

    def _imagem(par, png: bytes):
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.add_run().add_picture(io.BytesIO(png), width=largura_conteudo)

    if img_cab:
        _imagem(secao.header.paragraphs[0], img_cab)
    elif branding.get("cabecalho"):
        par = secao.header.paragraphs[0]
        par.text = branding["cabecalho"]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in par.runs:
            run.font.size = Pt(9)
            run.font.bold = True

    if img_rod:
        _imagem(secao.footer.paragraphs[0], img_rod)
    elif branding.get("rodape"):
        par = secao.footer.paragraphs[0]
        par.text = branding["rodape"]
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in par.runs:
            run.font.size = Pt(8)


def gerar_docx(titulo: str, texto_md: str, branding: dict | None = None) -> bytes:
    doc = _docx_novo()
    _docx_aplicar_branding(doc, branding)
    doc.add_paragraph(titulo.upper(), style="GovDocs Titulo")
    _docx_inserir_markdown(doc, texto_md)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def gerar_docx_consolidado(documentos: dict[str, str],
                           branding: dict | None = None,
                           dados: dict | None = None) -> bytes:
    doc = _docx_novo()
    _docx_aplicar_branding(doc, branding)
    doc.add_paragraph("DOCUMENTOS DA FASE PREPARATÓRIA — LEI Nº 14.133/2021",
                      style="GovDocs Titulo")
    doc.add_paragraph(f"Dossiê gerado em {date.today().strftime('%d/%m/%Y')}.",
                      style="GovDocs Nota")
    for doc_key in _ordem_de_exportacao(documentos, dados):
        doc.add_page_break()
        doc.add_paragraph(DOCUMENTOS[doc_key]["titulo"].upper(),
                          style="GovDocs Titulo")
        _docx_inserir_markdown(doc, documentos[doc_key])
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# PDF — caminho principal: DOCX estilizado -> LibreOffice -> PDF
# ---------------------------------------------------------------------------
_motor_pdf_efetivo: str | None = None


def motor_pdf() -> str:
    """
    Motor que REALMENTE será usado: 'libreoffice' (DOCX convertido —
    padrão institucional fiel) ou 'fpdf2' (renderizador próprio).

    Antes bastava o binário existir no PATH para o sistema anunciar
    'libreoffice' na tela de auditoria. Quando a conversão falhava em
    tempo de execução — o que acontece neste ambiente — o dossiê saía
    pelo fpdf2 e a interface seguia informando o motor errado, jogando
    para o vazio qualquer diagnóstico de formatação. Agora a resposta
    vem de uma conversão real de teste (feita uma vez por processo).
    """
    global _motor_pdf_efetivo
    if _motor_pdf_efetivo is not None:
        return _motor_pdf_efetivo
    if not (shutil.which("soffice") or shutil.which("libreoffice")):
        _motor_pdf_efetivo = "fpdf2"
        return _motor_pdf_efetivo
    try:
        sonda = _docx_novo()
        sonda.add_paragraph("sonda")
        buffer = io.BytesIO()
        sonda.save(buffer)
        convertido = _docx_em_pdf(buffer.getvalue())
    except Exception:
        convertido = None
    _motor_pdf_efetivo = "libreoffice" if convertido else "fpdf2"
    return _motor_pdf_efetivo


def _docx_em_pdf(docx_bytes: bytes) -> bytes | None:
    """Converte DOCX em PDF com o LibreOffice; None se indisponível/falhar."""
    import os
    import subprocess
    import tempfile

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        entrada = os.path.join(tmp, "documento.docx")
        with open(entrada, "wb") as fh:
            fh.write(docx_bytes)
        try:
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", tmp, entrada],
                check=True, capture_output=True, timeout=120,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired):
            return None
        saida = os.path.join(tmp, "documento.pdf")
        if not os.path.exists(saida):
            return None
        with open(saida, "rb") as fh:
            return fh.read()


def _pdf_aplicar_marca(pdf_bytes: bytes, branding: dict | None) -> bytes:
    """Marca d'água (imagem translúcida ou texto) SOB o texto, via PyMuPDF."""
    b = branding or {}
    img_marca = _img_bytes(b, "marca_img")
    texto_marca = (b.get("marca_dagua") or "").strip()
    if not img_marca and not texto_marca:
        return pdf_bytes
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for pagina in doc:
            w, h = pagina.rect.width, pagina.rect.height
            if img_marca:
                lado = w * 0.70
                rect = fitz.Rect((w - lado) / 2, h * 0.30,
                                 (w + lado) / 2, h * 0.30 + lado * 0.6)
                pagina.insert_image(rect, stream=img_marca, overlay=False)
            else:
                pagina.insert_text(
                    fitz.Point(w * 0.18, h * 0.60), texto_marca,
                    fontsize=48, rotate=90, color=(0.90, 0.90, 0.90),
                    overlay=False,
                )
        return doc.tobytes()
    except Exception:  # noqa: BLE001 — marca é acessório; nunca quebra o PDF
        return pdf_bytes


# ---------------------------------------------------------------------------
# PDF — fallback fpdf2 (sem LibreOffice no ambiente)
# ---------------------------------------------------------------------------
# As fontes nativas do fpdf2 usam Latin-1; substituímos os caracteres
# tipográficos comuns fora dessa tabela para não quebrar a exportação.
_SUBSTITUICOES_LATIN1 = {
    "—": "-", "–": "-", "‘": "'", "’": "'",
    "“": '"', "”": '"', "…": "...", "•": "-",
    " ": " ", "→": "->",
}


def _latin1_seguro(texto: str) -> str:
    for orig, subst in _SUBSTITUICOES_LATIN1.items():
        texto = texto.replace(orig, subst)
    return texto.encode("latin-1", errors="replace").decode("latin-1")


def _img_bytes(branding: dict, chave_b64: str, chave_legada: str = "") -> bytes | None:
    """Extrai PNG (bytes) de um campo base64 do branding, se houver."""
    from .branding import de_base64

    return de_base64((branding or {}).get(chave_b64) or "")


def _pdf_novo(branding: dict | None = None):
    """
    PDF A4 com identidade visual opcional do órgão.

    Prioriza IMAGENS capturadas de um documento-modelo (cabeçalho e rodapé
    carimbados na mesma posição relativa; marca d'água central translúcida).
    Sem imagens, cai para as versões em TEXTO (cabecalho/rodape/marca_dagua).
    branding pode conter: cabecalho_img, rodape_img, marca_img (base64 PNG),
    cabecalho_pct, rodape_pct (alturas em % da página) e os campos de texto.
    """
    from fpdf import FPDF

    b = branding or {}
    img_cab = _img_bytes(b, "cabecalho_img")
    img_rod = _img_bytes(b, "rodape_img")
    img_marca = _img_bytes(b, "marca_img")
    cab_pct = float(b.get("cabecalho_pct") or 14)
    rod_pct = float(b.get("rodape_pct") or 10)

    marca_txt = _latin1_seguro(b.get("marca_dagua") or "")
    cab_txt = _latin1_seguro(b.get("cabecalho") or "")
    rod_txt = _latin1_seguro(b.get("rodape") or "")

    # Alturas reservadas (mm) quando há imagem de cabeçalho/rodapé
    cab_h = (297.0 * cab_pct / 100.0) if img_cab else 0.0
    rod_h = (297.0 * rod_pct / 100.0) if img_rod else 0.0

    class PDFInstitucional(FPDF):
        def header(self):
            # Marca d'água (imagem central translúcida ou texto diagonal)
            if img_marca:
                larg = self.w * 0.7
                self.image(io.BytesIO(img_marca), x=(self.w - larg) / 2,
                           y=self.h * 0.28, w=larg)
            elif marca_txt:
                self.set_font("Helvetica", "B", 46)
                self.set_text_color(228, 228, 228)
                with self.rotation(45, self.w / 2, self.h / 2):
                    self.text(self.w / 2 - self.get_string_width(marca_txt) / 2,
                              self.h / 2, marca_txt)
                self.set_text_color(0, 0, 0)
            # Cabeçalho (imagem no topo, largura total) ou texto
            if img_cab:
                self.image(io.BytesIO(img_cab), x=0, y=0, w=self.w, h=cab_h)
            elif cab_txt:
                self.set_font("Helvetica", "B", 9)
                self.set_text_color(90, 90, 90)
                self.cell(0, 5, cab_txt, align="C", new_x="LMARGIN", new_y="NEXT")
                self.set_draw_color(200, 200, 200)
                self.line(self.l_margin, self.get_y() + 1,
                          self.w - self.r_margin, self.get_y() + 1)
                self.set_text_color(0, 0, 0)
            self.set_y(max(self.t_margin, cab_h + 4))

        def footer(self):
            if img_rod:
                self.image(io.BytesIO(img_rod), x=0, y=self.h - rod_h,
                           w=self.w, h=rod_h)
                return
            self.set_y(-14)
            self.set_font("Helvetica", "", 8)
            self.set_text_color(120, 120, 120)
            if rod_txt:
                self.cell(0, 4, rod_txt, align="C", new_x="LMARGIN", new_y="NEXT")
            self.cell(0, 4, f"Página {self.page_no()}/{{nb}}", align="C")
            self.set_text_color(0, 0, 0)

    pdf = PDFInstitucional(format="A4")
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=max(22, rod_h + 6))
    pdf.set_margins(left=20, top=max(20, cab_h + 6), right=20)
    return pdf


def _ponto_de_retorno(pdf):
    """
    Fotografa o PDF para que uma tentativa fracassada possa ser desfeita.

    Usa o `FPDFRecorder` do próprio fpdf2 — o mesmo mecanismo por trás de
    `FPDF.unbreakable()`. Ele guarda uma cópia profunda do estado (páginas
    já emitidas, buffer de conteúdo, cursor) e `rewind()` a restaura.

    Devolve `None` se o mecanismo não estiver disponível na versão
    instalada. Nesse caso quem chama NÃO tenta renderizar sem rede: vai
    direto ao caminho degradado, que é seguro. Uma tabela em parágrafos é
    feia; uma tabela com linhas repetidas é um ato administrativo que
    afirma item que o processo não tem.
    """
    try:
        from fpdf.recorder import FPDFRecorder
    except ImportError:      # pragma: no cover - fpdf2 sem o recorder
        return None
    try:
        return FPDFRecorder(pdf)
    except Exception:        # pragma: no cover - estado não copiável
        return None


def _pdf_render_tabela(pdf, linhas_tab: list[str]) -> None:
    """
    Renderiza uma tabela Markdown como tabela real do fpdf2 (com links).

    O fpdf2 NÃO divide uma linha entre páginas: uma célula muito longa
    (ex.: descrição de item vinda de planilha) estoura a altura da página
    e levanta ValueError ("row ... too high"). Estratégia: tenta fontes
    decrescentes; em último caso, degrada para parágrafos "Rótulo: valor"
    — o download nunca pode quebrar.
    """
    com_cabecalho = _tem_cabecalho(linhas_tab)
    linhas = [
        [_latin1_seguro(c) for c in _celulas_da_linha(ln)]
        for ln in linhas_tab
        if not _RE_SEPARADOR_TABELA.match(ln)  # descarta a linha ---|---
    ]
    if not linhas:
        return
    n = max(len(l) for l in linhas)
    linhas = [l + [""] * (n - len(l)) for l in linhas]
    linhas = _remover_colunas_vazias(linhas)
    n = len(linhas[0])
    largura = pdf.w - pdf.l_margin - pdf.r_margin
    # Colunas proporcionais ao texto RENDERIZADO: sem isto a Descrição do
    # item recebe a mesma fatia de 'Unidade' e o texto quebra a cada
    # poucos caracteres ("ESPECIFICA ÇÃO", "PASTA SANFONAD A").
    pesos = _pesos_das_colunas(linhas)

    for fonte_pt, altura_linha in ((9, 5), (7, 3.5), (6, 3)):
        # Cada tentativa é uma TRANSAÇÃO: ou a tabela inteira entra, ou o
        # PDF volta a ser exatamente o que era antes dela. O fpdf2 levanta
        # o ValueError NO MEIO do laço de linhas, com as anteriores já
        # desenhadas; sem desfazer, a tentativa seguinte as redesenhava e
        # o dossiê saía com item repetido — 20 linhas duas vezes com uma
        # queda, quatro vezes com três. Restaurar a posição do cursor não
        # resolvia: apaga o sintoma, não o conteúdo já emitido.
        retorno = _ponto_de_retorno(pdf)
        if retorno is None:
            break        # sem como desfazer, vai direto ao caminho seguro
        try:
            pdf.set_x(pdf.l_margin)
            pdf.set_font("Times", "", fonte_pt)
            with pdf.table(markdown=True,
                           first_row_as_headings=com_cabecalho,
                           col_widths=pesos,
                           line_height=altura_linha, width=largura) as tabela:
                for linha in linhas:
                    fpdf_linha = tabela.row()
                    for celula in linha:
                        fpdf_linha.cell(celula)
            # Com colunas de larguras diferentes, o cursor não retorna à
            # margem ao fim da tabela: o parágrafo seguinte começaria na
            # borda direita e sairia cortado da página.
            pdf.set_x(pdf.l_margin)
            pdf.ln(2)
            return
        except ValueError:
            # linha alta demais até para esta fonte — desfaz e tenta menor
            retorno.rewind()

    # Último recurso: conteúdo em parágrafos (nunca perde dados nem quebra)
    pdf.set_x(pdf.l_margin)
    cabecalho = linhas[0] if com_cabecalho else [""] * n
    pdf.set_font("Times", "", 10)
    for linha in (linhas[1:] if com_cabecalho else linhas):
        texto = "; ".join(
            f"{cab}: {val}" for cab, val in zip(cabecalho, linha) if val
        )
        pdf.multi_cell(largura, 5, _latin1_seguro(texto),
                       new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)
    pdf.ln(2)


def _pdf_inserir_markdown(pdf, texto_md: str) -> None:
    largura = pdf.w - pdf.l_margin - pdf.r_margin
    buffer_tab: list[str] = []

    def flush_tabela():
        if buffer_tab:
            _pdf_render_tabela(pdf, buffer_tab)
            buffer_tab.clear()

    for linha in texto_md.splitlines():
        tipo, conteudo = _classificar_linha(linha)
        if tipo == "tabela":
            buffer_tab.append(conteudo)
            continue
        flush_tabela()
        # títulos: sem marcação inline; corpo/itens: markdown (negrito e
        # links [texto](url) clicáveis e compactos)
        limpo = _latin1_seguro(_limpar_inline(conteudo))
        rico = _latin1_seguro(conteudo)
        # new_x/new_y explícitos: por padrão o fpdf2 deixa o cursor à
        # DIREITA da célula, na mesma linha. Dois parágrafos seguidos sem
        # linha em branco entre eles (1.1. e 1.2.) faziam o segundo
        # começar na borda direita e sair cortado da página — eram 203
        # blocos fora da margem no dossiê auditado.
        quebra = {"new_x": "LMARGIN", "new_y": "NEXT"}
        if tipo == "vazio":
            pdf.ln(3)
        elif tipo == "h1":
            pdf.set_font("Times", "B", 13)
            pdf.multi_cell(largura, 7, limpo, **quebra)
            pdf.ln(1)
        elif tipo == "h2":
            pdf.set_font("Times", "B", 12)
            pdf.multi_cell(largura, 6, limpo, **quebra)
            pdf.ln(1)
        elif tipo == "h3":
            pdf.set_font("Times", "B", 12)
            pdf.multi_cell(largura, 6, limpo, **quebra)
        elif tipo == "item":
            pdf.set_font("Times", "", 12)
            pdf.multi_cell(largura, 6.5, "  -  " + rico, markdown=True,
                           **quebra)
        else:
            pdf.set_font("Times", "", 12)
            pdf.multi_cell(largura, 6.5, rico, markdown=True, **quebra)
    flush_tabela()


def _pdf_bytes(pdf) -> bytes:
    saida = pdf.output()
    return bytes(saida)


def gerar_pdf(titulo: str, texto_md: str, branding: dict | None = None) -> bytes:
    """
    PDF do documento. Caminho principal: DOCX estilizado -> LibreOffice
    (mesmo conteúdo/formatação do DOCX, Times 12/1,5/6pt/justificado).
    Fallback: renderizador fpdf2 (fonte Times nativa).
    """
    convertido = _docx_em_pdf(gerar_docx(titulo, texto_md, branding))
    if convertido:
        return _pdf_aplicar_marca(convertido, branding)

    pdf = _pdf_novo(branding)
    pdf.add_page()
    pdf.set_font("Times", "B", 14)
    pdf.multi_cell(pdf.w - pdf.l_margin - pdf.r_margin, 8,
                   _latin1_seguro(titulo.upper()), align="C",
                   new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    _pdf_inserir_markdown(pdf, texto_md)
    return _pdf_bytes(pdf)


def gerar_pdf_consolidado(documentos: dict[str, str],
                          branding: dict | None = None,
                          dados: dict | None = None) -> bytes:
    convertido = _docx_em_pdf(
        gerar_docx_consolidado(documentos, branding, dados))
    if convertido:
        return _pdf_aplicar_marca(convertido, branding)

    pdf = _pdf_novo(branding)
    pdf.add_page()
    pdf.set_font("Times", "B", 14)
    largura = pdf.w - pdf.l_margin - pdf.r_margin
    pdf.multi_cell(
        largura, 8,
        _latin1_seguro("DOCUMENTOS DA FASE PREPARATÓRIA - LEI Nº 14.133/2021"),
        align="C", new_x="LMARGIN", new_y="NEXT",
    )
    pdf.set_font("Times", "", 10)
    pdf.multi_cell(largura, 6, f"Dossiê gerado em {date.today().strftime('%d/%m/%Y')}.",
                   new_x="LMARGIN", new_y="NEXT")
    for doc_key in _ordem_de_exportacao(documentos, dados):
        pdf.add_page()
        pdf.set_font("Times", "B", 13)
        pdf.multi_cell(largura, 8, _latin1_seguro(DOCUMENTOS[doc_key]["titulo"].upper()),
                       align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)
        _pdf_inserir_markdown(pdf, documentos[doc_key])
    return _pdf_bytes(pdf)


# ---------------------------------------------------------------------------
# Pacote ZIP com todos os arquivos individuais
# ---------------------------------------------------------------------------
def gerar_zip(documentos: dict[str, str], formato: str,
              branding: dict | None = None,
              dados: dict | None = None) -> bytes:
    """`formato`: 'docx' ou 'pdf'. Zipa um arquivo por documento aprovado."""
    def gerador(titulo, texto):
        fn = gerar_docx if formato == "docx" else gerar_pdf
        return fn(titulo, texto, branding)
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for doc_key in _ordem_de_exportacao(documentos, dados):
            # O número do arquivo é POSICIONAL na ordem canônica do
            # dossiê, não sequencial: o Edital é sempre 04, ainda que o
            # pacote saia incompleto.
            i = DOCUMENTOS_EXPORTAVEIS.index(doc_key) + 1
            meta = DOCUMENTOS[doc_key]
            nome = f"{i:02d}-{meta['sigla'].replace('/', '-')}.{formato}"
            zf.writestr(nome, gerador(meta["titulo"], documentos[doc_key]))
    return buffer.getvalue()
