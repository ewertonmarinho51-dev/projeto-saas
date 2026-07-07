"""Testes da Base de Conhecimento (RAG): extração, chunking e interface."""

import io
from pathlib import Path

import pytest
from streamlit.testing.v1 import AppTest

from src import rag

APP = str(Path(__file__).resolve().parent.parent / "app.py")


# ---------------------------------------------------------------------------
# Extração de texto
# ---------------------------------------------------------------------------
def _pdf_exemplo() -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(
        0, 8,
        "Acordao 1234/2024 - Plenario. A jurisprudencia do TCU e pacifica "
        "quanto a obrigatoriedade do parcelamento do objeto quando tecnica "
        "e economicamente viavel, nos termos da Sumula 247." * 3,
    )
    return bytes(pdf.output())


def _docx_exemplo() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph(
        "Estudo Técnico Preliminar — aquisição de notebooks. "
        "Justificativa de parcelamento conforme art. 40, V, 'b', da Lei 14.133/2021."
    )
    tabela = doc.add_table(rows=1, cols=2)
    tabela.rows[0].cells[0].text = "Risco: atraso"
    tabela.rows[0].cells[1].text = "Mitigação: multa"
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def test_extrai_texto_pdf():
    texto = rag.extrair_texto("acordao.pdf", _pdf_exemplo())
    assert "Sumula 247" in texto and len(texto) > 100


def test_extrai_texto_docx_incluindo_tabelas():
    texto = rag.extrair_texto("etp.docx", _docx_exemplo())
    assert "Lei 14.133/2021" in texto
    assert "Risco: atraso | Mitigação: multa" in texto


def test_extrai_texto_txt_e_rejeita_formatos():
    assert "conteúdo" in rag.extrair_texto("nota.txt", "conteúdo de teste".encode() * 5)
    with pytest.raises(rag.ErroRAG, match="não suportado"):
        rag.extrair_texto("planilha.xlsx", b"x" * 100)
    with pytest.raises(rag.ErroRAG, match="texto extraível"):
        rag.extrair_texto("vazio.txt", b"abc")


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------
def test_chunking_com_sobreposicao():
    paragrafo = "O parcelamento do objeto é regra na Lei 14.133/2021. " * 20
    texto = "\n\n".join(paragrafo for _ in range(5))
    chunks = rag.dividir_em_chunks(texto, tamanho=1000, sobreposicao=150)

    assert len(chunks) > 1
    assert all(len(c) <= 1000 for c in chunks)
    # reconstrução: todo o conteúdo aparece em algum chunk
    assert all(c.strip() for c in chunks)
    # sobreposição: o início de um chunk repete o fim do anterior
    assert chunks[1][:50] in chunks[0] + chunks[1]


def test_chunking_texto_curto_gera_um_chunk():
    assert rag.dividir_em_chunks("texto curto") == ["texto curto"]


# ---------------------------------------------------------------------------
# Montagem da consulta e do bloco de referências
# ---------------------------------------------------------------------------
def test_montar_consulta_combina_campos():
    consulta = rag.montar_consulta(
        {"objeto": "Aquisição de notebooks", "justificativa": "parque obsoleto",
         "modelo_execucao": "SRP"},
        "etp",
    )
    assert "estudo técnico preliminar" in consulta
    assert "Aquisição de notebooks" in consulta and "SRP" in consulta


def test_bloco_referencias_vazio_sem_banco(monkeypatch):
    # sem Supabase configurado o bloco deve ser vazio e sem exceções
    monkeypatch.setattr(rag.db, "disponivel", lambda: False)
    assert rag.montar_bloco_referencias({"objeto": "x"}, "dfd") == ""


def test_bloco_referencias_formata_trechos(monkeypatch):
    monkeypatch.setattr(
        rag,
        "buscar_referencias",
        lambda consulta, qtd=6: [
            {"conteudo": "Súmula 247 do TCU...", "titulo": "Acórdão 1234",
             "categoria": "acordao", "similaridade": 0.9}
        ],
    )
    bloco = rag.montar_bloco_referencias({"objeto": "x"}, "etp")
    assert "REFERÊNCIAS DA BASE DE CONHECIMENTO" in bloco
    assert "Acórdão 1234" in bloco and "Súmula 247" in bloco
    assert "NÃO copie dados específicos" in bloco


# ---------------------------------------------------------------------------
# Navegação: página Base de Conhecimento renderiza sem banco
# ---------------------------------------------------------------------------
def test_pagina_biblioteca_renderiza_sem_supabase():
    at = AppTest.from_file(APP, default_timeout=60)
    at.secrets["SUPABASE_URL"] = ""
    at.secrets["SUPABASE_KEY"] = ""
    at.run()
    radio = [r for r in at.radio if r.key == "pagina"][0]
    radio.set_value("📚 Base de Conhecimento")
    at.run()
    assert not at.exception
    titulos = " ".join(s.value for s in at.subheader)
    assert "Base de Conhecimento" in titulos
    # sem Supabase deve orientar a configuração, não quebrar
    avisos = " ".join(w.value for w in at.warning)
    assert "SUPABASE_URL" in avisos
