"""
Testes obrigatórios da AUDITORIA P0 (plano de correção do sistema de
geração documental), construídos com trechos LITERAIS dos documentos
reais gerados pela Prefeitura de Paragominas:

  1.  isolamento de campos (cargo/função com valor improvisado);
  2.  prioridade (URL da planilha vazada na prosa);
  3.  tabela de itens (injeção única, em bloco próprio; cabeçalho DOCX
      apenas quando existe cabeçalho Markdown real; linha longa pode
      dividir entre páginas);
  4.  datas/estado (reiniciar_processo não vaza estado entre contratações);
  5.  placeholders (matrícula provisória 999999/15 sinalizada);
  6.  fundamentos legais (pregão × art. 109; vigência da ata × art. 82;
      pagamento × art. 98);
  7.  repactuação × reajuste em contratação de bens/materiais;
  8.  consistência do dossiê (tabela duplicada por documento);
  9.  garantia contratual "seca" (só percentual);
  10. CNPJ com dígitos verificadores inválidos.
"""

import io
import re

import streamlit as st

from src import achados, planilha, state, validacao
from src.export import gerar_docx


def _bloqueios(doc_key: str, texto: str) -> list[str]:
    return [a["mensagem"]
            for a in validacao.bloqueios(validacao.validar_documento(doc_key, texto))]


def _avisos(doc_key: str, texto: str) -> list[str]:
    return [a["mensagem"]
            for a in validacao.avisos(validacao.validar_documento(doc_key, texto))]


# ---------------------------------------------------------------------------
# 1. Isolamento de campos: cargo/função não aceita número/escala improvisada
# ---------------------------------------------------------------------------
def test_isolamento_campos_cargo_com_numero_bloqueia():
    # trecho literal do DFD real gerado
    texto = ("9.1.2. Representante da área de almoxarifado/controle de "
             "estoque: 15.\n")
    assert any("função/cargo preenchido com valor inválido" in m
               for m in _bloqueios("dfd", texto))


def test_isolamento_campos_cargo_com_escala_bloqueia():
    # trecho literal do ETP real gerado
    texto = ("16.1.2. Representante da área de almoxarifado/controle de "
             "estoque: alto.\n")
    assert any("função/cargo preenchido com valor inválido" in m
               for m in _bloqueios("etp", texto))


def test_isolamento_campos_cargo_com_nome_real_passa():
    texto = ("9.1.1. Responsável pela formalização da demanda: "
             "Luan Jardel de Moura Santos, Diretor de Compras.\n")
    assert not any("função/cargo" in m for m in _bloqueios("dfd", texto))


# ---------------------------------------------------------------------------
# 2. Prioridade: URL da planilha não pode vazar para a prosa
# ---------------------------------------------------------------------------
def test_prioridade_com_url_da_planilha_bloqueia():
    # trecho literal do DFD real gerado (item da planilha colado no campo)
    texto = ("1.5. Prioridade da demanda: 3(https://www.tkshopping.com.br/"
             "produto/alfinete-mapa-colorido-sortido-50un-pct-com-10) |\n")
    assert any("URL crua na prosa" in m for m in _bloqueios("dfd", texto))


def test_url_institucional_na_prosa_nao_bloqueia():
    texto = ("As propostas serão recebidas exclusivamente no sistema "
             "disponível em www.gov.br/compras, na forma do edital.\n")
    assert not any("URL crua" in m for m in _bloqueios("edital", texto))


def test_link_compacto_em_linha_de_tabela_nao_bloqueia():
    texto = ("| Código | Descrição | Fonte / Link |\n|---|---|---|\n"
             "| 001 | Caneta azul | [link](https://www.loja.com/caneta) |\n")
    assert not any("URL crua" in m for m in _bloqueios("dfd", texto))


# ---------------------------------------------------------------------------
# 3. Tabela de itens: injeção única, em bloco próprio; export robusto
# ---------------------------------------------------------------------------
_ITENS = [
    {"codigo": f"{i:03d}", "descricao": f"Item de expediente número {i} "
     "com descrição detalhada de especificação técnica para teste",
     "unidade": "un", "quantidade": 2, "valor_unitario": 10.0}
    for i in range(1, 15)  # > LIMITE_ITENS_INLINE: fluxo com marcador
]


def test_marcador_no_meio_da_frase_vira_bloco_proprio():
    texto = ("## 7. ESTIMATIVA DE VALOR\n\nA planilha consolidada, a ser "
             "juntada ao processo conforme marcação "
             f"{planilha.MARCADOR_TABELA} e demais anexos.")
    resultado = planilha.injetar_tabela(texto, _ITENS)
    # a linha de cabeçalho abre uma linha própria (nunca colada na prosa)
    assert "conforme marcação | Código" not in resultado
    assert "\n\n| Código | Descrição |" in resultado


def test_marcador_duplicado_injeta_tabela_uma_unica_vez():
    texto = (f"## 1. GERAL\n{planilha.MARCADOR_TABELA}\n\n"
             f"## 7. VALOR\n{planilha.MARCADOR_TABELA}\n")
    resultado = planilha.injetar_tabela(texto, _ITENS)
    assert len(re.findall(r"^\|\s*Código\s*\|", resultado, re.M)) == 1
    assert planilha.MARCADOR_TABELA not in resultado


def _tabelas_docx(docx_bytes: bytes):
    from docx import Document

    return Document(io.BytesIO(docx_bytes)).tables


def test_docx_sem_separador_nao_promove_item_a_cabecalho():
    # cabeçalho Markdown colado na prosa: a 1ª linha da tabela é DADO e
    # não pode virar cabeçalho repetido em toda página (defeito real:
    # item 572704 repetido 95 vezes no DFD)
    md = ("Texto conforme marcação | Código | Descrição |\n"
          "|---|---|\n"
          "| 572704 | ALFINETE PARA MURAL |\n"
          "| 572705 | ALMOFADA PARA CARIMBO |\n")
    tabela = _tabelas_docx(gerar_docx("DFD", md))[0]
    xml = tabela.rows[0]._tr.xml
    assert "tblHeader" not in xml
    assert not any(r.bold for p in tabela.rows[0].cells[0].paragraphs
                   for r in p.runs)


def test_docx_com_cabecalho_real_mantem_repeticao_de_cabecalho():
    md = ("| Código | Descrição |\n|---|---|\n| 001 | Caneta |\n")
    tabela = _tabelas_docx(gerar_docx("DFD", md))[0]
    assert "tblHeader" in tabela.rows[0]._tr.xml


def test_docx_linha_longa_pode_dividir_entre_paginas():
    descricao = "ESPECIFICAÇÃO detalhada do item com texto muito longo " * 10
    md = (f"| Código | Descrição |\n|---|---|\n| 001 | {descricao} |\n"
          "| 002 | Curta |\n")
    tabela = _tabelas_docx(gerar_docx("DFD", md))[0]
    assert "cantSplit" not in tabela.rows[1]._tr.xml   # longa: pode dividir
    assert "cantSplit" in tabela.rows[2]._tr.xml       # curta: não divide


# ---------------------------------------------------------------------------
# 4. Datas/estado: nada do processo anterior vaza para a próxima contratação
# ---------------------------------------------------------------------------
def test_reiniciar_processo_limpa_estado_transitorio():
    st.session_state["dados"] = {"prazo": "até março/2026"}
    st.session_state["documentos"] = {"dfd": "texto"}
    st.session_state["aprovados"] = {"dfd"}
    st.session_state["processo_id"] = "abc"
    st.session_state["_ciclo_resultado"] = {"hash": "x"}
    st.session_state["_ciclo_manual"] = True
    st.session_state["_fatos_cache"] = {"chave": "y"}
    st.session_state["_familia_escolha_dfd"] = "familia-1"
    st.session_state["_memorando_lido"] = "file-1"
    st.session_state["registro_geracoes"] = [{"doc": "dfd"}]
    # estado GLOBAL da sessão: sobrevive ao reinício do processo
    st.session_state["usuario"] = {"id": "u1", "nome": "Ana"}
    st.session_state["tenant_id"] = "t1"
    st.session_state["api_key_manual"] = "chave"
    st.session_state["_modelo_img"] = b"png-preview-do-admin"
    try:
        state.reiniciar_processo()
    except Exception:  # noqa: BLE001 — st.rerun() exige o runtime do Streamlit
        pass           # a limpeza ocorre antes do rerun
    assert st.session_state["dados"] == {}
    assert st.session_state["documentos"] == {}
    assert st.session_state["processo_id"] is None
    for chave in ("_ciclo_resultado", "_ciclo_manual", "_fatos_cache",
                  "_familia_escolha_dfd", "_memorando_lido",
                  "registro_geracoes"):
        assert chave not in st.session_state, chave
    # autenticação, tenant, chaves de API e caches globais preservados
    assert st.session_state["usuario"]["nome"] == "Ana"
    assert st.session_state["tenant_id"] == "t1"
    assert st.session_state["api_key_manual"] == "chave"
    assert st.session_state["_modelo_img"] == b"png-preview-do-admin"
    # higiene: session_state é global no modo bare — remove o que o teste
    # criou para não contaminar os demais testes do processo
    for chave in ("usuario", "tenant_id", "api_key_manual", "_modelo_img",
                  "dados", "documentos", "aprovados", "processo_id", "etapa"):
        st.session_state.pop(chave, None)


# ---------------------------------------------------------------------------
# 5. Placeholders: matrícula provisória é sinalizada; [PREENCHER] segue
#    bloqueando (regra pré-existente)
# ---------------------------------------------------------------------------
def test_matricula_999999_gera_aviso():
    # trecho literal do DFD real gerado
    texto = ("1.3. Responsável pela formalização da demanda: Luan Jardel "
             "de Moura Santos — matrícula:999999.\n")
    assert any("matrícula com aparência de improviso" in m
               for m in _avisos("dfd", texto))


def test_matricula_15_gera_aviso_e_matricula_real_passa():
    assert any("matrícula com aparência" in m
               for m in _avisos("dfd", "Servidor João — matrícula: 15.\n"))
    assert not any("matrícula" in m
                   for m in _avisos("dfd", "Servidora Ana — matrícula: 48291.\n"))


def test_preencher_continua_bloqueando():
    assert any("campo pendente" in m
               for m in _bloqueios("dfd", "Prazo: [PREENCHER: prazo em dias].\n"))


# ---------------------------------------------------------------------------
# 6. Fundamentos legais (RAG/citações): confusões recorrentes interceptadas
# ---------------------------------------------------------------------------
def test_pregao_fundado_no_art_109_bloqueia():
    # trecho literal do Edital real gerado
    texto = ("1.3. Modalidade: Pregão Eletrônico, na forma do art. 109 e "
             "seguintes da Lei nº 14.133/2021 e do edital.\n")
    assert any("pregão fundamentado no art. 109" in m
               for m in _bloqueios("edital", texto))


def test_pregao_com_fundamento_correto_passa():
    texto = ("1.3. Modalidade: Pregão Eletrônico, na forma dos arts. 28, I, "
             "e 29 da Lei nº 14.133/2021.\n")
    assert not any("fundamento legal" in m for m in _bloqueios("edital", texto))


def test_vigencia_da_ata_no_art_82_gera_aviso_e_art_84_passa():
    # trecho literal do Edital real gerado
    errado = ("2.5. Prazo de vigência da Ata de Registro de Preços: 12 "
              "(doze) meses, contado da publicação do extrato da Ata "
              "(art. 82, Lei nº 14.133/2021).\n")
    certo = ("2.5. Prazo de vigência da Ata de Registro de Preços: 12 "
             "(doze) meses, prorrogável por igual período (art. 84 da "
             "Lei nº 14.133/2021).\n")
    assert any("vigência da Ata" in m for m in _avisos("edital", errado))
    assert not any("vigência da Ata" in m for m in _avisos("edital", certo))


def test_pagamento_fundado_no_art_98_gera_aviso():
    # trecho literal do ETP real gerado
    texto = ("Vedação de pagamento antecipado sem entrega (art. 98 e "
             "art. 103).\n")
    assert any("pagamento fundado no art. 98" in m
               for m in _avisos("etp", texto))


# ---------------------------------------------------------------------------
# 7. Repactuação × reajuste em contratação de bens/materiais
# ---------------------------------------------------------------------------
def test_repactuacao_sem_mao_de_obra_gera_aviso():
    # trecho literal do TR real gerado (materiais de expediente)
    texto = ("3.10. Prazo de vigência e repactuação: vigência recomendada "
             "de 12 meses com previsão de critérios objetivos para "
             "repactuação de preços somente nas hipóteses legais.\n")
    assert any("repactuação" in m for m in _avisos("tr", texto))


def test_repactuacao_com_mao_de_obra_passa():
    texto = ("Serviço contínuo com dedicação exclusiva de mão de obra; "
             "admite-se repactuação nos termos do art. 135 da Lei nº "
             "14.133/2021.\n")
    assert not any("repactuação" in m and "inadequado" in m
                   for m in _avisos("tr", texto))


# ---------------------------------------------------------------------------
# 8. Consistência do dossiê: tabela de itens única por documento
# ---------------------------------------------------------------------------
def test_tabela_de_itens_duplicada_no_mesmo_documento_bloqueia():
    tabela = ("| Código | Descrição | Unidade |\n|---|---|---|\n"
              "| 001 | Caneta | un |\n")
    doc = f"## 1. GERAL\n{tabela}\n## 7. VALOR\n{tabela}"
    assert any("tabela de itens duplicada" in m for m in _bloqueios("dfd", doc))


def test_tabela_unica_por_documento_no_dossie_passa():
    tabela = ("| Código | Descrição | Unidade |\n|---|---|---|\n"
              "| 001 | Caneta | un |\n")
    achados_v = validacao.validar_todos({
        "dfd": f"## 7. VALOR\n{tabela}",
        "tr": f"## 5. VALOR\n{tabela}",
    })
    assert not any("duplicada" in a["mensagem"] for a in achados_v)


# ---------------------------------------------------------------------------
# 9. Garantia contratual "seca" (só o percentual, sem condições)
# ---------------------------------------------------------------------------
def test_garantia_so_percentual_gera_aviso():
    # trecho literal da Ata (Edital real gerado)
    assert any("garantia sem fundamentação" in m
               for m in _avisos("edital", "9.2. Garantia contratual: 5%.\n"))


def test_garantia_desenvolvida_passa():
    texto = ("9.2. A garantia contratual será prestada na modalidade "
             "caução, seguro-garantia ou fiança bancária, no percentual "
             "de 5% do valor do contrato, nos termos dos arts. 96 a 98 "
             "da Lei nº 14.133/2021.\n")
    assert not any("garantia sem fundamentação" in m
                   for m in _avisos("edital", texto))


# ---------------------------------------------------------------------------
# 10. CNPJ: dígitos verificadores conferidos deterministicamente
# ---------------------------------------------------------------------------
def test_cnpj_invalido_bloqueia_e_valido_passa():
    invalido = "Contratada: Empresa X, CNPJ 12.345.678/0001-00.\n"
    valido = "Contratada: Empresa X, CNPJ 12.345.678/0001-95.\n"
    assert any("CNPJ inválido" in m for m in _bloqueios("edital", invalido))
    assert not any("CNPJ" in m for m in _bloqueios("edital", valido))


def test_cnpj_zerado_bloqueia():
    assert any("CNPJ inválido" in m
               for m in _bloqueios("edital",
                                   "CNPJ 00.000.000/0000-00 (a confirmar).\n"))


# ---------------------------------------------------------------------------
# Integração com o ciclo de correção: os novos achados viram findings
# ---------------------------------------------------------------------------
def test_url_crua_vira_finding_corrigivel():
    texto = ("## 2. JUSTIFICATIVA\n\nVer preço em "
             "https://www.tkshopping.com.br/produto/alfinete no anexo.\n")
    relatorio = achados.gerar_relatorio({"memo": texto})
    f = next(x for x in relatorio["findings"]
             if x["categoria"] == "vazamento_planilha")
    assert f["autoCorrectable"] is True
    assert f["severity"] == "HIGH"


def test_matricula_improvisada_vira_finding_de_dado_ausente():
    texto = "## 1. GERAL\n\nServidor João — matrícula: 999999.\n"
    relatorio = achados.gerar_relatorio({"memo": texto})
    f = next(x for x in relatorio["findings"]
             if x["categoria"] == "dado_improvisado")
    assert f["autoCorrectable"] is False
    assert f["blockingReason"] == achados.MOTIVO_DADO_AUSENTE
