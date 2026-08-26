"""
Ata de Registro de Preços obsoleta — conflito semântico apontado pela
auditoria independente da Fase 1.

A ARP **não** pertence a `SEQUENCIA_DOCUMENTOS` (não é etapa do wizard),
mas pertence a `DOCUMENTOS_EXPORTAVEIS`. Enquanto a invalidação de estado
percorria só a sequência, esta trajetória exportava uma Ata que já não
correspondia ao processo:

    SRP → gera Edital + ARP → volta ao formulário → muda para NÃO-SRP
    → DFD/ETP/TR/Edital são invalidados → a ARP antiga PERMANECE em
    `session_state.documentos` → tudo é regerado → o exportador percorre
    DOCUMENTOS_EXPORTAVEIS e leva junto a Ata obsoleta.

Duas linhas de defesa são provadas aqui: a limpeza de estado (a Ata cai
com aquilo que a fundamenta) e a decisão de exportação (processo sem SRP
não exporta ARP nem com uma chave 'arp' residual).
"""

import io
import os
import zipfile
from pathlib import Path

import pytest
import streamlit as st
from streamlit.testing.v1 import AppTest

from src import export, state
from src.config import (DOCUMENTOS_EXPORTAVEIS, SEQUENCIA_DOCUMENTOS,
                        adota_srp, exportaveis_do_processo)

APP = str(Path(__file__).resolve().parent.parent / "app.py")

COM_SRP = {"orgao": "Prefeitura", "objeto": "materiais de expediente",
           "modelo_execucao": "Sistema de Registro de Preços (SRP)"}
SEM_SRP = {**COM_SRP, "modelo_execucao": "Fornecimento integral"}

QUATRO = {k: f"## 1. OBJETO\n\nTexto do {k}.\n" for k in SEQUENCIA_DOCUMENTOS}
CINCO = {**QUATRO, "arp": "## 1. OBJETO\n\nAta de Registro de Preços.\n"}


@pytest.fixture
def sessao():
    """Estado de wizard limpo, com a ARP já gerada sob SRP."""
    st.session_state.clear()
    st.session_state["dados"] = dict(COM_SRP)
    st.session_state["documentos"] = dict(CINCO)
    st.session_state["edicoes_pendentes"] = {}
    st.session_state["aprovados"] = set(DOCUMENTOS_EXPORTAVEIS)
    st.session_state["etapa"] = 5
    yield st.session_state
    st.session_state.clear()


# ---------------------------------------------------------------------------
# 1ª linha de defesa — limpeza de estado
# ---------------------------------------------------------------------------
def test_alterar_o_formulario_leva_a_arp_junto(sessao):
    """Qualquer mudança upstream torna a Ata potencialmente obsoleta."""
    state.invalidar_a_partir_de("formulario")
    assert sessao["documentos"] == {}
    assert sessao["aprovados"] == set()


def test_regerar_o_edital_descarta_a_ata_que_saiu_com_ele(sessao):
    """
    A Ata é emitida JUNTO do edital. Descartar o edital e manter a Ata
    seria reaproveitamento silencioso — a Ata seguiria a modelagem
    anterior sem que ninguém tivesse decidido isso.
    """
    state.descartar_documento("edital")
    assert "edital" not in sessao["documentos"]
    assert "arp" not in sessao["documentos"]
    assert "arp" not in sessao["aprovados"]
    # os anteriores continuam: eles não dependem do edital
    assert set(sessao["documentos"]) == {"dfd", "etp", "tr"}


def test_editar_o_edital_aprovado_invalida_a_ata(sessao):
    """`invalidar_a_partir_de('edital')` alcança o instrumento derivado."""
    state.invalidar_a_partir_de("edital")
    assert "arp" not in sessao["documentos"]
    assert "edital" in sessao["documentos"]   # o próprio, não


def test_alterar_o_tr_derruba_edital_e_ata(sessao):
    state.invalidar_a_partir_de("tr")
    assert set(sessao["documentos"]) == {"dfd", "etp", "tr"}


# ---------------------------------------------------------------------------
# 2ª linha de defesa — a decisão de exportação
# ---------------------------------------------------------------------------
def test_processo_sem_srp_nao_exporta_arp_residual():
    """
    Mesmo que a limpeza de estado falhe, a chave 'arp' residual não pode
    virar arquivo: quem decide é o formulário do processo.
    """
    assert not adota_srp(SEM_SRP)
    assert exportaveis_do_processo(SEM_SRP, CINCO) == SEQUENCIA_DOCUMENTOS
    assert exportaveis_do_processo(COM_SRP, CINCO) == DOCUMENTOS_EXPORTAVEIS


def test_zip_sem_srp_ignora_a_arp_residual():
    nomes = zipfile.ZipFile(io.BytesIO(
        export.gerar_zip(CINCO, "pdf", None, SEM_SRP))).namelist()
    assert len(nomes) == 4, nomes
    assert not any("ARP" in n for n in nomes), nomes
    # e o número do arquivo continua POSICIONAL: o Edital é sempre 04
    assert "04-Edital.pdf" in nomes, nomes


def test_zip_com_srp_leva_a_arp():
    nomes = zipfile.ZipFile(io.BytesIO(
        export.gerar_zip(CINCO, "pdf", None, COM_SRP))).namelist()
    assert len(nomes) == 5, nomes
    assert "05-ARP.pdf" in nomes, nomes


def test_dossie_docx_sem_srp_nao_contem_a_ata():
    from docx import Document

    doc = Document(io.BytesIO(
        export.gerar_docx_consolidado(CINCO, None, SEM_SRP)))
    titulos = [p.text for p in doc.paragraphs]
    assert not any("ATA DE REGISTRO" in t.upper() for t in titulos), titulos

    doc = Document(io.BytesIO(
        export.gerar_docx_consolidado(CINCO, None, COM_SRP)))
    titulos = [p.text for p in doc.paragraphs]
    assert any("ATA DE REGISTRO" in t.upper() for t in titulos), titulos


def test_sem_o_formulario_o_exportador_nao_inventa_resposta():
    """
    Minuta de trabalho e bundle avulso não trazem o formulário. Aí não há
    como saber se a Ata cabe — exporta-se o que veio, em vez de adivinhar.
    """
    nomes = zipfile.ZipFile(io.BytesIO(
        export.gerar_zip(CINCO, "pdf"))).namelist()
    assert len(nomes) == 5, nomes


# ---------------------------------------------------------------------------
# As duas transições exigidas, ponta a ponta no estado
# ---------------------------------------------------------------------------
def test_transicao_srp_para_nao_srp_nao_deixa_ata_para_tras(sessao):
    """SRP → Edital+ARP → formulário → NÃO-SRP → regenera → 4 documentos."""
    assert state.exportaveis() == DOCUMENTOS_EXPORTAVEIS

    # o usuário volta ao formulário e muda a modelagem
    state.invalidar_a_partir_de("formulario")
    sessao["dados"] = dict(SEM_SRP)

    # e regera o processo inteiro — sem SRP, a Ata não é produzida
    for chave in SEQUENCIA_DOCUMENTOS:
        sessao["documentos"][chave] = QUATRO[chave]
        sessao["aprovados"].add(chave)

    assert "arp" not in sessao["documentos"]
    assert state.exportaveis() == SEQUENCIA_DOCUMENTOS
    nomes = zipfile.ZipFile(io.BytesIO(export.gerar_zip(
        sessao["documentos"], "pdf", None, sessao["dados"]))).namelist()
    assert not any("ARP" in n for n in nomes), nomes


def test_transicao_nao_srp_para_srp_produz_e_exporta_a_ata(sessao):
    """Caminho inverso: a Ata nasce quando a modelagem passa a ser SRP."""
    sessao["dados"] = dict(SEM_SRP)
    sessao["documentos"] = dict(QUATRO)
    sessao["aprovados"] = set(SEQUENCIA_DOCUMENTOS)
    assert state.exportaveis() == SEQUENCIA_DOCUMENTOS

    state.invalidar_a_partir_de("formulario")
    sessao["dados"] = dict(COM_SRP)
    for chave in SEQUENCIA_DOCUMENTOS:
        sessao["documentos"][chave] = QUATRO[chave]
        sessao["aprovados"].add(chave)
    # a geração do edital emite a Ata junto (steps.render_documento)
    sessao["documentos"]["arp"] = CINCO["arp"]

    assert state.exportaveis() == DOCUMENTOS_EXPORTAVEIS
    nomes = zipfile.ZipFile(io.BytesIO(export.gerar_zip(
        sessao["documentos"], "pdf", None, sessao["dados"]))).namelist()
    assert "05-ARP.pdf" in nomes, nomes


# ---------------------------------------------------------------------------
# A tela final não oferece o que não pode ser emitido
# ---------------------------------------------------------------------------
def _tela_final(dados, documentos):
    os.environ["GOVDOCS_MODO_ABERTO"] = "1"
    at = AppTest.from_file(APP, default_timeout=60)
    at.secrets["SUPABASE_URL"] = ""
    at.secrets["SUPABASE_KEY"] = ""
    at.session_state["etapa"] = 5
    at.session_state["dados"] = dict(dados)
    at.session_state["documentos"] = dict(documentos)
    at.session_state["aprovados"] = set(documentos)
    return at


def test_tela_final_sem_srp_nao_oferece_download_da_ata():
    at = _tela_final(SEM_SRP, CINCO)
    at.run()
    assert not at.exception
    rotulos = [b.label or "" for b in at.get("download_button")]
    assert not any("ARP" in r for r in rotulos), rotulos
    assert any("ZIP com os 4 PDFs" in r for r in rotulos), rotulos


def test_tela_final_com_srp_oferece_a_ata():
    at = _tela_final(COM_SRP, CINCO)
    at.run()
    assert not at.exception
    rotulos = [b.label or "" for b in at.get("download_button")]
    assert any("ARP" in r for r in rotulos), rotulos
    assert any("ZIP com os 5 PDFs" in r for r in rotulos), rotulos
