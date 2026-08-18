"""Estilos institucionais do DOCX e motor de PDF (padrão dos docs manuais)."""

import io
import zipfile

from docx import Document

from src import export

MD = (
    "## 1. DO OBJETO\n\n"
    "1.1. Aquisição de material de expediente para a Administração.\n\n"
    "1.1.1. Subitem de terceiro nível com detalhamento.\n\n"
    "| Código | Descrição | Valor |\n|---|---|---|\n| 1 | Caneta | R$ 1,00 |\n\n"
    "________________________\nResponsável pela demanda\n"
)


# Margem institucional de 2 cm, em pontos (72 pt = 1 polegada). É a
# margem da própria página, não um número escolhido até o teste passar:
# fora dela ficam cabeçalho e rodapé, dentro dela fica o CORPO — a única
# região sobre a qual a regra tipográfica institucional fala.
MARGEM_EM_PONTOS = 2 / 2.54 * 72


def _doc(md: str = MD) -> Document:
    return Document(io.BytesIO(export.gerar_docx("Termo de Referência", md)))


def test_corpo_times_12_15_6pt_justificado():
    doc = _doc()
    normal = doc.styles["Normal"]
    assert normal.font.name == "Times New Roman"
    assert normal.font.size.pt == 12
    corpo = doc.styles["GovDocs Corpo"]
    pf = corpo.paragraph_format
    assert pf.line_spacing == 1.5
    assert pf.space_after.pt == 6
    assert str(pf.alignment) == "JUSTIFY (3)"
    assert pf.widow_control is True


def test_clausula_negrito_presa_ao_conteudo():
    doc = _doc()
    estilo = doc.styles["GovDocs Clausula"]
    assert estilo.font.bold is True
    assert estilo.paragraph_format.keep_with_next is True
    clausulas = [p for p in doc.paragraphs if p.style.name == "GovDocs Clausula"]
    assert clausulas and "DO OBJETO" in clausulas[0].text


def test_itens_hierarquicos_com_recuo():
    doc = _doc()
    estilos = [p.style.name for p in doc.paragraphs]
    assert "GovDocs Item 1" in estilos   # 1.1.
    assert "GovDocs Item 2" in estilos   # 1.1.1.


def test_assinatura_estilo_proprio():
    doc = _doc()
    assinaturas = [p for p in doc.paragraphs if p.style.name == "GovDocs Assinatura"]
    assert assinaturas


def test_margens_a4():
    doc = _doc()
    s = doc.sections[0]
    assert round(s.page_width.cm, 1) == 21.0
    assert round(s.top_margin.cm, 1) == 2.5
    assert round(s.left_margin.cm, 1) == 2.0


def test_tabela_com_cabecalho_repetido_e_linha_indivisivel():
    docx = export.gerar_docx("TR", MD)
    with zipfile.ZipFile(io.BytesIO(docx)) as zf:
        xml = zf.read("word/document.xml").decode()
    assert "tblHeader" in xml     # cabeçalho repete nas páginas seguintes
    assert "cantSplit" in xml     # linha não divide entre páginas


def test_pdf_via_libreoffice_quando_disponivel():
    pdf = export.gerar_pdf("Termo de Referência", MD)
    assert pdf.startswith(b"%PDF")
    if export.motor_pdf() == "libreoffice":
        # O CORPO usa a família serifada (Times/Liberation/Nimbus) — que
        # é o que o comentário sempre disse e a asserção não dizia.
        #
        # A versão anterior recusava Helvetica em QUALQUER posição da
        # página e falhava por causa do rodapé "Página 1/1", que é
        # sans-serif de propósito. Falhava, portanto, por motivo
        # diferente do declarado — e teste assim treina quem lê a suíte
        # a ignorá-lo.
        import fitz

        doc = fitz.open(stream=pdf, filetype="pdf")
        pagina = doc[0]
        limite_do_rodape = pagina.rect.height - MARGEM_EM_PONTOS
        spans = [s for b in pagina.get_text("dict")["blocks"]
                 for l in b.get("lines", []) for s in l.get("spans", [])]
        corpo = [s for s in spans
                 if MARGEM_EM_PONTOS < s["bbox"][1] < limite_do_rodape]
        assert corpo, "nenhum texto no corpo da página"

        def _serifada(fonte: str) -> bool:
            return ("Times" in fonte or "Liberation" in fonte
                    or "Nimbus" in fonte)

        fontes = {s["font"] for s in corpo}
        # mais forte que a asserção antiga: TODO o corpo é serifado, e
        # não apenas "existe alguma serifada em algum lugar da página"
        assert all(_serifada(f) for f in fontes), (
            f"corpo com fonte não serifada: {fontes}")

        # E o rodapé é FIXADO, não apenas ignorado: excluir uma região
        # sem dizer o que se espera dela transforma a exclusão numa
        # gaveta onde qualquer regressão futura cabe.
        rodape = [s for s in spans if s["bbox"][1] >= limite_do_rodape]
        assert rodape, "rodapé sumiu do PDF"
        assert all("Página" in s["text"] or _serifada(s["font"])
                   for s in rodape), [s["text"] for s in rodape]


def test_pdf_consolidado_mesmo_conteudo_do_docx():
    docs = {"dfd": MD}
    pdf = export.gerar_pdf_consolidado(docs)
    assert pdf.startswith(b"%PDF")
    import fitz

    texto = "".join(p.get_text() for p in fitz.open(stream=pdf, filetype="pdf"))
    assert "DO OBJETO" in texto and "Caneta" in texto


def test_fallback_fpdf2_linha_gigante_nao_estoura(monkeypatch):
    """Regressão do crash em produção: fpdf2 'row too high' com descrição
    enorme — deve degradar (fonte menor/parágrafos), nunca levantar."""
    monkeypatch.setattr(export, "_docx_em_pdf", lambda b: None)  # força fallback
    descricao = "ESPECIFICAÇÃO: " + "texto muito longo de item, " * 150
    md = ("## 5. ESTIMATIVA\n\n| Código | Descrição | Qtd |\n|---|---|---|\n"
          f"| 001 | {descricao} | 100 |\n")
    pdf = export.gerar_pdf_consolidado({"dfd": md})
    assert pdf.startswith(b"%PDF")
