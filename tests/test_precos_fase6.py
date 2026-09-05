"""
Fase 6 — relatórios, exportações e as duas correções que ela expôs.

O relatório completo é a memória do ato administrativo. As provas aqui
guardam três coisas que, se cederem, tornam o documento inútil ou
enganoso:

1. ele contém o que foi **descartado**, com motivo (§31.14);
2. ele **não inventa** — ausência é escrita, não deixada em branco;
3. o identificador da versão identifica o **resultado**, não a emissão.

E duas correções em `export.py`, ambas descobertas ao rodar o pipeline
de verdade com dados realistas: um caractere fora do latin-1 numa célula
derrubava a geração inteira, e preencher tabela por `Table.cell` era
quadrático no tamanho da tabela.
"""

from __future__ import annotations

import io
import zipfile
from datetime import datetime, timezone

import pytest

from src import export
from src.precos import relatorio

pytestmark = pytest.mark.usefixtures("motor_institucional")


PESQUISA = {
    "id": "pesq-abc12345", "versao": 2, "nome": "Material de expediente",
    "estado": "completed", "perfil_normativo": "in_65_2021",
    "data_base": "2026-09-01", "objeto": "Aquisição de material",
    "responsavel": "Servidora Fulana", "local_referencia": "Paragominas/PA",
    "processo_id": "proc-1", "criado_em": "2026-08-15",
    "motivo_da_revisao": "troca de metodologia",
    "filtros": {"janela_dias": 365, "uf": "PA"},
}


def _item(numero=1, **extras) -> dict:
    base = {
        "id": f"i{numero}", "numero": numero,
        "descricao": f"ITEM {numero} — MATERIAL DE EXPEDIENTE",
        "unidade": "UNIDADE", "quantidade": "100", "estado": "complete",
        "metodo": "mediana", "preco_estimado": "2.35",
        "preco_total": "235.00", "codigo": "236168",
        "tipo_catalogo": "CATMAT",
        "justificativa": "mediana de 4 referências\nteto da mediana aplicado",
        "estatisticas": {
            "estatisticas": {"quantidade": 4, "menor": "2.00",
                             "maior": "9.00", "media": "3.50",
                             "mediana": "2.35",
                             "coeficiente_variacao": "0.55"},
            "anomalias": [{"valor": "9.00", "criterio": "IQR",
                           "motivo": "283% acima da mediana"}],
        },
    }
    base.update(extras)
    return base


def _ref(sufixo="1", status="selected", **extras) -> dict:
    base = {
        "id": f"r{sufixo}", "status": status, "fonte_nome": "Compras.gov",
        "fonte_id": "compras_gov_precos", "fonte_tipo": "sistema_oficial",
        "orgao": "PREFEITURA MUNICIPAL", "uf": "PA",
        "data_resultado": "2026-06-01", "quantidade_original": "100",
        "unidade_normalizada": "UNIDADE",
        "valor_unitario_normalizado": "2.00", "score": "0.88",
        "identidade": "0.95", "circunstancias": "0.93",
        "id_externo": f"e{sufixo}", "raw_hash": "a" * 64, "motivos": [],
    }
    base.update(extras)
    return base


CENARIO = ([_item(1)], {"i1": [
    _ref("1"),
    _ref("2", status="rejected", fonte_nome="PNCP", fonte_id="pncp",
         fonte_tipo="contratacao_similar", uf="SP",
         unidade_original="CAIXA", unidade_normalizada=None,
         valor_unitario_normalizado=None, valor_unitario_original="40.00",
         score="0.31", raw_hash="b" * 64,
         motivos=["unidade da referência é embalagem e a fonte não "
                  "informou quantos itens ela contém"]),
]})


# ===========================================================================
# §31 — relatório completo
# ===========================================================================
# Os itens 13 a 17 do §31 são POR ITEM (referências usadas, descartadas,
# memória, unitário, total): existe um conjunto deles para cada item, e
# promovê-los a seções de topo os faria aparecer uma vez só. Ficam dentro
# do bloco de cada item, e o relatório diz isso em voz alta.
SECOES_DE_TOPO = [n for n in range(1, 23) if n not in range(13, 18)]


def test_as_vinte_e_duas_secoes_estao_presentes():
    itens, refs = CENARIO
    texto = relatorio.completo(PESQUISA, itens, refs)
    for n in SECOES_DE_TOPO:
        assert f"## {n}." in texto, f"seção {n} ausente"
    # e os itens 13–17 aparecem, por item, com a correspondência explícita
    assert "item 13 da estrutura" in texto
    assert "Referências selecionadas" in texto
    assert "Referências desconsideradas" in texto
    assert "Memória de cálculo" in texto
    assert "Preço estimado unitário (§31.16)" in texto
    assert "Valor total do item (§31.17)" in texto


def test_o_descartado_aparece_com_o_motivo():
    """
    §31.14 — a pergunta que a auditoria faz é justamente por que os
    outros preços não entraram. Um relatório que mostrasse só a cesta
    seria uma defesa, não uma memória.
    """
    itens, refs = CENARIO
    texto = relatorio.completo(PESQUISA, itens, refs)
    assert "Referências desconsideradas (1)" in texto
    assert "embalagem" in texto
    assert "excluída" in texto
    # e a selecionada continua listada à parte
    assert "Referências selecionadas (1)" in texto


def test_o_relatorio_nao_inventa_o_que_nao_existe():
    """
    Campo em branco é dúvida sobre se ninguém preencheu ou se o sistema
    perdeu. "(não informado)" é informação.
    """
    magro = {"id": "p", "versao": 1, "nome": "Sem nada"}
    texto = relatorio.completo(magro, [_item(1, codigo=None,
                                             tipo_catalogo=None)], {"i1": []})
    assert relatorio.AUSENTE in texto
    assert "a busca foi feita por descrição" in texto


def test_a_ressalva_nao_vira_conclusao_juridica():
    """
    §23: dispersão estatística não é inexequibilidade nem irregularidade,
    e o relatório vai para processo administrativo.
    """
    itens, refs = CENARIO
    texto = relatorio.completo(PESQUISA, itens, refs).lower()
    assert "sinalizado para revisão" in texto
    # O relatório NEGA a conclusão jurídica explicitamente; o que não
    # pode é AFIRMÁ-LA. Por isso a prova procura as formas assertivas, e
    # não a palavra solta — que aparece justamente na negação.
    assert "não afirma inexequibilidade nem irregularidade" in texto
    for afirmacao in ("preço inexequível", "preco inexequivel",
                      "preço ilegal", "é irregular", "houve "
                      "superfaturamento", "indício de sobrepreço"):
        assert afirmacao not in texto, afirmacao


def test_o_perfil_normativo_vai_escrito_no_relatorio():
    """
    §3: a IN 65 não é norma municipal automática. O relatório precisa
    dizer sob qual regra o valor foi formado.
    """
    texto = relatorio.completo(PESQUISA, *CENARIO)
    assert "IN SEGES/ME nº 65/2021" in texto
    assert "Teto da mediana" in texto


def test_itens_sem_preco_nao_entram_no_valor_global():
    itens = [_item(1), _item(2, estado="incomplete", preco_total=None,
                             preco_estimado=None)]
    refs = {"i1": [_ref("1")], "i2": []}
    texto = relatorio.completo(PESQUISA, itens, refs)
    assert "R$ 235,00" in texto
    assert "**não** os inclui" in texto
    assert "**Itens sem preço formado:** 2." in texto


# ===========================================================================
# §34 — identificador da versão
# ===========================================================================
def test_o_identificador_e_do_resultado_e_nao_da_emissao():
    """
    Dois relatórios do mesmo resultado, emitidos em dias diferentes, têm
    de bater — senão o identificador não serve para provar nada.
    """
    itens, refs = CENARIO
    ontem = datetime(2026, 9, 4, tzinfo=timezone.utc)
    hoje = datetime(2026, 9, 5, tzinfo=timezone.utc)
    a = relatorio.completo(PESQUISA, itens, refs, emitido_em=ontem)
    b = relatorio.completo(PESQUISA, itens, refs, emitido_em=hoje)
    assert a != b                      # a data de emissão muda o texto…
    identificador = relatorio.identificador_da_versao(PESQUISA, itens, refs)
    assert identificador in a and identificador in b   # …o identificador não


@pytest.mark.parametrize("mudanca", [
    {"preco_estimado": "9.99"},
    {"metodo": "media"},
    {"estado": "incomplete"},
])
def test_mudar_o_resultado_muda_o_identificador(mudanca):
    itens, refs = CENARIO
    antes = relatorio.identificador_da_versao(PESQUISA, itens, refs)
    depois = relatorio.identificador_da_versao(
        PESQUISA, [_item(1, **mudanca)], refs)
    assert antes != depois


def test_reclassificar_uma_referencia_muda_o_identificador():
    """
    Excluir uma referência da cesta muda o resultado, ainda que o preço
    final não mude. O identificador precisa refletir isso.
    """
    itens, refs = CENARIO
    antes = relatorio.identificador_da_versao(PESQUISA, itens, refs)
    outras = {"i1": [_ref("1", status="rejected"), refs["i1"][1]]}
    assert relatorio.identificador_da_versao(PESQUISA, itens, outras) != antes


def test_o_identificador_nao_depende_da_ordem_das_referencias():
    itens, refs = CENARIO
    invertidas = {"i1": list(reversed(refs["i1"]))}
    assert (relatorio.identificador_da_versao(PESQUISA, itens, refs)
            == relatorio.identificador_da_versao(PESQUISA, itens, invertidas))


# ===========================================================================
# §32 — quadro resumido
# ===========================================================================
def test_o_quadro_traz_os_tres_primeiros_precos():
    """
    Três não é número mágico: é a regra dos três preços, e mostrá-los
    permite conferir o cálculo de cabeça.
    """
    refs = {"i1": [_ref(str(n), valor_unitario_normalizado=f"{n}.00")
                   for n in range(1, 6)]}
    texto = relatorio.resumido(PESQUISA, [_item(1)], refs)
    assert "Preço 1" in texto and "Preço 3" in texto
    assert "R$ 1,00" in texto and "R$ 3,00" in texto
    assert "| 2 |" in texto            # duas "outras" referências


def test_o_quadro_usa_travessao_e_nao_texto_longo():
    """
    Numa tabela de 210 linhas, "(não informado)" repetido em três
    colunas alarga a tabela até estourar a largura útil da página — o
    defeito de geometria que a Fase 2.1 fechou.
    """
    texto = relatorio.resumido(PESQUISA, [_item(1)], {"i1": []})
    linha = [ln for ln in texto.splitlines() if ln.startswith("| 01 ")][0]
    assert relatorio.VAZIO_NO_QUADRO in linha
    assert relatorio.AUSENTE not in linha


def test_a_concordancia_esta_certa():
    """
    "dos 1 item(ns)" denuncia texto de máquina e tira credibilidade do
    que está certo no resto da página.
    """
    um = relatorio.resumido(PESQUISA, [_item(1)], {"i1": []})
    assert "de 1 item concluído" in um
    dois = relatorio.resumido(PESQUISA, [_item(1), _item(2)],
                              {"i1": [], "i2": []})
    assert "dos 2 itens concluídos" in dois


# ===========================================================================
# §33 — memória analítica
# ===========================================================================
def test_o_xlsx_tem_as_tres_abas_e_inclui_o_descartado():
    from openpyxl import load_workbook

    itens, refs = CENARIO
    wb = load_workbook(io.BytesIO(
        relatorio.xlsx_analitico(PESQUISA, itens, refs)))
    assert wb.sheetnames == ["Itens", "Referências", "Identificação"]

    detalhe = wb["Referências"]
    situacoes = [linha[1] for linha in detalhe.iter_rows(min_row=2,
                                                         values_only=True)]
    assert "selecionada" in situacoes and "excluída" in situacoes


def test_ausencia_de_preco_vira_celula_vazia_e_nao_zero():
    """
    Zero numa planilha de preços é um preço; vazio é a ausência dele. A
    diferença muda a soma que alguém vai fazer na coluna.
    """
    from openpyxl import load_workbook

    itens, refs = CENARIO
    wb = load_workbook(io.BytesIO(
        relatorio.xlsx_analitico(PESQUISA, itens, refs)))
    detalhe = wb["Referências"]
    # a referência excluída não tem preço normalizado
    linha = [l for l in detalhe.iter_rows(min_row=2, values_only=True)
             if l[1] == "excluída"][0]
    assert linha[15] is None, linha[15]


def test_o_identificador_da_versao_esta_na_planilha():
    from openpyxl import load_workbook

    itens, refs = CENARIO
    wb = load_workbook(io.BytesIO(
        relatorio.xlsx_analitico(PESQUISA, itens, refs)))
    ficha = {l[0]: l[1] for l in wb["Identificação"].iter_rows(
        min_row=2, values_only=True)}
    assert ficha["Identificador da versão"] == \
        relatorio.identificador_da_versao(PESQUISA, itens, refs)


def test_o_nome_do_arquivo_e_estavel_e_legivel():
    nome = relatorio.nome_do_arquivo(PESQUISA, "completo", "pdf")
    assert nome == "Material-de-expediente-rev2-completo.pdf"


# ===========================================================================
# §33 — pelo motor institucional, sem segundo pipeline
# ===========================================================================
def test_o_pdf_sai_pelo_motor_institucional():
    itens, refs = CENARIO
    markdown = relatorio.completo(PESQUISA, itens, refs)
    pdf = export.gerar_pdf("Relatório de Pesquisa de Preços", markdown)
    assert pdf.startswith(b"%PDF")
    assert export.motor_pdf() == "libreoffice"


def test_o_docx_sai_com_as_tabelas_do_relatorio():
    from docx import Document

    itens, refs = CENARIO
    markdown = relatorio.completo(PESQUISA, itens, refs)
    doc = Document(io.BytesIO(
        export.gerar_docx("Relatório", markdown)))
    assert len(doc.tables) >= 3        # fontes, quadro-resumo, referências
    texto = "\n".join(p.text for p in doc.paragraphs)
    assert "Identificador da versão" in texto


def test_o_pacote_traz_os_tres_arquivos():
    from src.ui import precos_ui

    itens, refs = CENARIO
    pacote = precos_ui.montar_pacote(PESQUISA, itens, refs)
    with zipfile.ZipFile(io.BytesIO(pacote)) as zf:
        nomes = sorted(zf.namelist())
    assert len(nomes) == 3
    assert any(n.endswith("-completo.pdf") for n in nomes)
    assert any(n.endswith("-resumido.pdf") for n in nomes)
    assert any(n.endswith("-memoria-analitica.xlsx") for n in nomes)


# ===========================================================================
# As duas correções em export.py
# ===========================================================================
@pytest.mark.parametrize("celula", [
    "CANETA — AZUL",           # travessão: vem de todo texto colado do Word
    "CANETA – AZUL",           # meia-risca
    "CANETA “AZUL”",           # aspas curvas
    "CANETA … AZUL",           # reticências
    "CANETA 🖊 AZUL",           # fora até do mapa de equivalentes
])
def test_caractere_fora_do_latin1_na_tabela_nao_derruba_a_exportacao(celula):
    """
    O defeito que a Fase 6 expôs, e que atingia o PRODUTO INTEIRO.

    O medidor de largura usa Times em latin-1. Travessão, meia-risca e
    aspas curvas — presentes em qualquer descrição colada do Word ou
    extraída de PDF — levantavam `FPDFUnicodeEncodingException` e
    derrubavam a geração do DOCX e do PDF do processo, não só a do
    relatório de preços.

    Medir é cálculo auxiliar: não pode ser o que impede o documento de
    existir.
    """
    markdown = ("# TESTE\n\n| Código | Descrição | Valor |\n"
                f"|---|---|---|\n| 001 | {celula} | R$ 2,35 |\n")
    docx = export.gerar_docx("Teste", markdown)
    assert docx and len(docx) > 1000
    assert export.gerar_pdf("Teste", markdown).startswith(b"%PDF")


def test_a_largura_do_travessao_e_estimada_e_nao_ignorada():
    """
    Superestimar deixa a coluna um pouco larga; subestimar estoura a
    página. O equivalente de medida não pode ser vazio.
    """
    largura = export._largura_de_texto_cm("A—B")      # noqa: SLF001
    assert largura > export._largura_de_texto_cm("AB")  # noqa: SLF001


def test_a_tabela_nao_e_preenchida_celula_a_celula(monkeypatch):
    """
    A correção de desempenho, medida de forma DETERMINÍSTICA.

    No python-docx, `Table.cell` reconstrói a lista de todas as células
    da tabela a cada acesso — preencher célula a célula é quadrático no
    tamanho da tabela. Com um relatório de 15 itens, `cell()` respondia
    por 14,4 s de 18,5 s.

    Contar as chamadas é melhor que cronometrar: não depende da máquina
    e aponta a causa, não o sintoma.
    """
    from docx.table import Table

    chamadas = {"n": 0}
    original = Table.cell

    def contando(self, *args, **kwargs):
        chamadas["n"] += 1
        return original(self, *args, **kwargs)

    monkeypatch.setattr(Table, "cell", contando)

    linhas = ["| A | B | C | D |", "|---|---|---|---|"]
    linhas += [f"| {n} | x | y | z |" for n in range(60)]
    export.gerar_docx("Teste", "# T\n\n" + "\n".join(linhas))

    assert chamadas["n"] == 0, (
        f"Table.cell foi chamado {chamadas['n']} vezes — o preenchimento "
        "voltou a ser quadrático")
